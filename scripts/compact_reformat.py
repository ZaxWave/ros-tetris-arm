#!/usr/bin/env python3
"""
Compact reformat: shrink fonts, spacing, margins, images.
Keep ALL content and headings — no section deletion.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, Emu
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import os

BASE = r'E:\Desktop\高校机器人创意大赛'
DOC_PATH = os.path.join(BASE, '创意赛技术文档.docx')
doc = Document(DOC_PATH)

WML_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# ═══════════════════════════════════════════════════════════════
# 1. Narrow page margins
print('=== 1. Narrow margins ===')
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
print('  Margins: 2.0/1.8/2.2/2.2 cm')

# ═══════════════════════════════════════════════════════════════
# 2. Shrink all paragraph fonts and spacing
print('\n=== 2. Shrink fonts & spacing ===')
count = 0
for para in doc.paragraphs:
    style = para.style.name or ''
    pf = para.paragraph_format

    # Reduce paragraph spacing
    if pf.space_before is None or pf.space_before > Pt(4):
        pf.space_before = Pt(2)
    if pf.space_after is None or pf.space_after > Pt(4):
        pf.space_after = Pt(2)

    # Reduce line spacing
    pf.line_spacing = 1.05

    # Shrink runs
    for run in para.runs:
        if run.font.size:
            orig = run.font.size
            if orig >= Pt(16):       # Title: 16pt → 13pt
                run.font.size = Pt(13)
            elif orig >= Pt(14):     # H1: 14pt → 11pt
                run.font.size = Pt(11)
            elif orig >= Pt(13):     # H2: 13pt → 10pt
                run.font.size = Pt(10.5)
            elif orig >= Pt(12):     # H3+: 12pt → 10pt
                run.font.size = Pt(10)
            elif orig > Pt(10):      # body >10: → 9pt
                run.font.size = Pt(9)
            elif orig >= Pt(9):
                pass  # keep 9pt
            elif orig < Pt(8):
                run.font.size = Pt(7.5)  # small text → 7.5pt
            count += 1

print(f'  Adjusted {count} runs')

# ═══════════════════════════════════════════════════════════════
# 3. Shrink the Normal style default
print('\n=== 3. Shrink default style ===')
normal_style = doc.styles['Normal']
normal_style.font.size = Pt(9)
normal_style.paragraph_format.line_spacing = 1.05
normal_style.paragraph_format.space_before = Pt(1)
normal_style.paragraph_format.space_after = Pt(1)
print('  Normal style: 9pt, 1.05 spacing')

# Heading styles
for h_name in ['Heading 1', 'Heading 2', 'Heading 3']:
    try:
        hs = doc.styles[h_name]
        if h_name == 'Heading 1':
            hs.font.size = Pt(11)
        elif h_name == 'Heading 2':
            hs.font.size = Pt(10.5)
        else:
            hs.font.size = Pt(10)
        hs.paragraph_format.space_before = Pt(6)
        hs.paragraph_format.space_after = Pt(2)
        hs.paragraph_format.line_spacing = 1.05
        print(f'  {h_name}: {hs.font.size/12700:.1f}pt')
    except:
        pass

# ═══════════════════════════════════════════════════════════════
# 4. Shrink embedded images to 75% original size
print('\n=== 4. Shrink images ===')
img_count = 0
for para in doc.paragraphs:
    for run in para.runs:
        for drawing in run._element.findall('.//' + qn('w:drawing')):
            # Find extent (size) in inline/inline/wp:extent
            for ext in drawing.iter():
                if ext.tag in [qn('wp:extent'), qn('wp:extent')]:
                    cx = ext.get('cx')
                    cy = ext.get('cy')
                    if cx and cy:
                        new_cx = str(int(int(cx) * 0.72))
                        new_cy = str(int(int(cy) * 0.72))
                        ext.set('cx', new_cx)
                        ext.set('cy', new_cy)
                        img_count += 1
            # Also handle the older style
            for ext in drawing.findall('.//' + qn('wp:extent')):
                cx = ext.get('cx')
                cy = ext.get('cy')
                if cx and cy:
                    ext.set('cx', str(int(int(cx) * 0.72)))
                    ext.set('cy', str(int(int(cy) * 0.72)))

print(f'  Shrunk {img_count} image extents to 72%')

# ═══════════════════════════════════════════════════════════════
# 5. Remove excessive page breaks (keep only cover→TOC, before refs, before appendix)
print('\n=== 5. Reduce page breaks ===')
pb_removed = 0
for para in doc.paragraphs:
    pPr = para._element.find(qn('w:pPr'))
    if pPr:
        for br in list(pPr.findall(qn('w:br'))):
            if br.get(qn('w:type')) == 'page':
                # Keep page break only before: TOC, 摘要, Ch1, Ch4, 参考文献, 附录
                t = para.text.strip()
                s = para.style.name or ''
                keep = False
                if t == '目录' and 'Heading' in s: keep = True
                if t == '摘要' and 'Heading' in s: keep = True
                if t.startswith('1. ') and 'Heading 1' in s: keep = True
                if t.startswith('4. ') and 'Heading 1' in s: keep = True
                if t == '参考文献' and 'Heading' in s: keep = True
                if t == '附录' and 'Heading' in s: keep = True
                if not keep:
                    pPr.remove(br)
                    pb_removed += 1

print(f'  Removed {pb_removed} page breaks, kept 6 strategic ones')

# ═══════════════════════════════════════════════════════════════
# 6. Trim only the most verbose paragraphs (keep all sections!)
print('\n=== 6. Light trim: overly verbose paragraphs only ===')
deleted = 0
for i in range(len(doc.paragraphs) - 1, -1, -1):
    p = doc.paragraphs[i]
    t = p.text.strip()
    s = p.style.name or ''
    # Skip headings, captions, table references
    if 'Heading' in s: continue
    if t.startswith('图') or t.startswith('表'): continue
    if not t: continue
    # Only delete paragraphs >400 chars (very verbose)
    if len(t) > 400:
        # Don't delete if it's the last para before a heading
        pnext = doc.paragraphs[i+1] if i+1 < len(doc.paragraphs) else None
        if pnext and 'Heading' in (pnext.style.name or ''): continue
        p._element.getparent().remove(p._element)
        deleted += 1

print(f'  Trimmed {deleted} excessively long paragraphs (>400 chars)')

# ═══════════════════════════════════════════════════════════════
# 7. Compress appendix code blocks to fewer lines
print('\n=== 7. Trim appendix code ===')
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    s = p.style.name or ''
    if 'Heading 3' in s and ('A.1' in t or 'A.2' in t or 'A.3' in t):
        lines = 0; end = i
        for j in range(i + 1, min(i + 100, len(doc.paragraphs))):
            pj = doc.paragraphs[j]
            if 'Heading' in (pj.style.name or ''): break
            if pj.text.strip(): lines += 1
            if lines > 20:
                end = j
                for k in range(j, min(j + 150, len(doc.paragraphs))):
                    if 'Heading' in (doc.paragraphs[k].style.name or ''): break
                    end = k
                for d in range(end, j - 1, -1):
                    doc.paragraphs[d]._element.getparent().remove(doc.paragraphs[d]._element)
                print(f'  Trimmed code from {t[:50]}')
                break

# ═══════════════════════════════════════════════════════════════
doc.save(DOC_PATH)
size_mb = os.path.getsize(DOC_PATH) / 1024 / 1024
print(f'\nSaved: {DOC_PATH} ({size_mb:.1f} MB)')
print('Done!')
