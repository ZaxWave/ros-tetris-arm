#!/usr/bin/env python3
"""Surgically trim document from 54 to ≤30 pages."""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import os

BASE = r'E:\Desktop\高校机器人创意大赛'
DOC_PATH = os.path.join(BASE, '创意赛技术文档.docx')

doc = Document(DOC_PATH)


def find_heading(text_pattern, style='Heading 2'):
    """Find paragraph index of a heading containing text."""
    for i, para in enumerate(doc.paragraphs):
        if style in (para.style.name or '') and text_pattern in para.text:
            return i
    return None


def find_heading1(text_pattern):
    return find_heading(text_pattern, 'Heading 1')


def find_heading2(text_pattern):
    return find_heading(text_pattern, 'Heading 2')


def delete_range(start, end):
    """Delete paragraphs from start to end (inclusive). Returns count deleted."""
    count = 0
    for i in range(end, start - 1, -1):
        doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
        count += 1
    return count


def delete_section(heading_idx, include_subsections=True, stop_before=None):
    """Delete from a heading to the next heading of same or higher level."""
    if heading_idx is None:
        return 0
    heading_level = int(doc.paragraphs[heading_idx].style.name.split()[-1])
    end = len(doc.paragraphs) - 1
    for i in range(heading_idx + 1, len(doc.paragraphs)):
        s = doc.paragraphs[i].style.name or ''
        if 'Heading' in s:
            level = int(s.split()[-1])
            if level <= heading_level:
                end = i - 1
                break
        if stop_before and stop_before in doc.paragraphs[i].text:
            end = i - 1
            break
    return delete_range(heading_idx, end)


# ═══════════════════════════════════════════════════════════════
print('=== 1. Remove Appendix A (code listings) ===')
# Keep only a reference note, remove all code
app_a = find_heading2('附录A')
if app_a:
    # Find 附录B to stop before
    app_b = find_heading2('附录B')
    n = delete_range(app_a + 1, app_b - 1)  # Delete content but keep heading
    # Change heading to a brief note
    doc.paragraphs[app_a].text = ''  # clear then set via runs
    for run in doc.paragraphs[app_a].runs:
        run.text = ''
    doc.paragraphs[app_a].runs[0].text = '附录A  核心代码说明'
    # Add a single line referencing GitHub
    note_p = doc.add_paragraph()
    note_p.add_run('核心控制代码（yolo_detector.py、moveJ_P.py、tetris_planner.py）已开源，完整源码见 GitHub 仓库：').font.size = Pt(9)
    note_p2 = doc.add_paragraph()
    r = note_p2.add_run('https://github.com/ZaxWave/ros-tetris-arm')
    r.font.size = Pt(9); r.font.name = 'Consolas'
    # Move to right after heading
    body = doc.element.body
    note_p2._element.getparent().remove(note_p2._element)
    note_p._element.getparent().remove(note_p._element)
    ref = doc.paragraphs[app_a]._element
    ref.getparent().insert(list(ref.getparent()).index(ref) + 1, note_p2._element)
    ref.getparent().insert(list(ref.getparent()).index(ref) + 1, note_p._element)
    print(f'  Replaced Appendix A ({n} paragraphs) with GitHub reference')


# ═══════════════════════════════════════════════════════════════
print('\n=== 2. Remove Appendix D (FAQ) ===')
app_d = find_heading2('附录D')
if app_d:
    n = delete_section(app_d)
    print(f'  Deleted Appendix D: {n} paragraphs')


# ═══════════════════════════════════════════════════════════════
print('\n=== 3. Remove duplicate 附图C-3 ===')
# Find 附图C-3 and its associated image paragraph
for i, para in enumerate(doc.paragraphs):
    if '附图C-3' in para.text and 'YOLO' in para.text:
        # Delete: image paragraph before it (P i-1) + this caption (P i)
        prev = doc.paragraphs[i-1]
        para._element.getparent().remove(para._element)
        prev._element.getparent().remove(prev._element)
        print(f'  Deleted 附图C-3 (image P{i-1} + caption P{i})')
        break


# ═══════════════════════════════════════════════════════════════
print('\n=== 4. Compress Chapter 3: Dev & Ops → brief note ===')
ch4_idx = find_heading1('4. ')
# Find Ch3 heading
ch3_idx = find_heading1('3. ')
ch3_end = ch4_idx - 1 if ch4_idx else None

if ch3_idx and ch3_end:
    # Keep the heading and one summary paragraph
    # Delete all content between heading+1 and ch3_end
    summary_text = (
        '系统开发环境基于 Ubuntu 20.04 + ROS Noetic，工作空间位于 ~/ws_rmrobot/。'
        '启动流程为：roscore → rm_driver → realsense2_camera 节点，'
        '随后通过 vision_pipeline.py 执行视觉识别，tetris_planner 进行拼接决策，'
        'moveJ_P.py 完成拾放任务执行。详细启动命令见 GitHub 仓库 README。'
    )
    n = delete_range(ch3_idx + 1, ch3_end)
    # Insert summary paragraph after heading
    summary_p = doc.add_paragraph()
    summary_p.add_run(summary_text).font.size = Pt(10)
    summary_p._element.getparent().remove(summary_p._element)
    ref = doc.paragraphs[ch3_idx]._element
    ref.getparent().insert(list(ref.getparent()).index(ref) + 1, summary_p._element)
    print(f'  Compressed Ch3: {n} paragraphs → 1 summary paragraph')


# ═══════════════════════════════════════════════════════════════
print('\n=== 5. Compress references ===')
refs_idx = find_heading1('参考文献')
app_idx = find_heading1('附录')
if refs_idx and app_idx:
    # Keep only the first 15 references (essential ones), delete rest
    ref_lines = 0
    delete_from = None
    for i in range(refs_idx + 1, app_idx):
        t = doc.paragraphs[i].text.strip()
        if t and t[0].isdigit() and '[' in t[:5]:
            ref_lines += 1
        if ref_lines > 15 and delete_from is None:
            delete_from = i
    if delete_from:
        n = delete_range(delete_from, app_idx - 1)
        print(f'  Trimmed references: kept 15, deleted {n} paragraphs')


# ═══════════════════════════════════════════════════════════════
print('\n=== 6. Compress Appendix B (experimental data) ===')
# Keep the table, remove verbose descriptions
app_b = find_heading2('附录B')
app_c = find_heading2('附录C')
if app_b and app_c:
    # Delete empty/verbose paragraphs between heading and table
    for i in range(app_b + 1, app_c):
        t = doc.paragraphs[i].text.strip()
        # Keep table captions (start with 附表B), delete verbose text
        if t and not t.startswith('附表B') and 'Heading' not in (doc.paragraphs[i].style.name or ''):
            # Check if paragraph has a table nearby — keep if very short
            if len(t) > 150:
                doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
                print(f'  Deleted verbose para in App B: {t[:60]}...')


# ═══════════════════════════════════════════════════════════════
print('\n=== 7. Trim verbose subsections in Ch2 ===')
# Remove subsection 2.1.3 (计算与通信系统) entirely if verbose
# Actually let's just reduce some of the very verbose paragraphs

ch2_idx = find_heading1('2. ')
ch3_idx = find_heading1('3. ')
if ch2_idx and ch3_idx:
    # Delete paragraphs longer than 300 chars in Ch2 (verbose descriptions)
    deleted = 0
    for i in range(ch3_idx - 1, ch2_idx, -1):
        t = doc.paragraphs[i].text.strip()
        style = doc.paragraphs[i].style.name or ''
        if 'Heading' not in style and len(t) > 300 and deleted < 4:
            doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
            deleted += 1
    print(f'  Trimmed {deleted} verbose paragraphs in Ch2')


# ═══════════════════════════════════════════════════════════════
print('\n=== 8. Trim Ch4 verbose paragraphs ===')
ch4_idx = find_heading1('4. ')
ch5_idx = find_heading1('5. ')
if ch4_idx and ch5_idx:
    deleted = 0
    for i in range(ch5_idx - 1, ch4_idx, -1):
        t = doc.paragraphs[i].text.strip()
        style = doc.paragraphs[i].style.name or ''
        if 'Heading' not in style and len(t) > 350 and deleted < 6:
            doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
            deleted += 1
    print(f'  Trimmed {deleted} verbose paragraphs in Ch4')


# ═══════════════════════════════════════════════════════════════
doc.save(DOC_PATH)
size_mb = os.path.getsize(DOC_PATH) / 1024 / 1024
print(f'\nSaved: {DOC_PATH} ({size_mb:.1f} MB)')
print('Done trimming!')
