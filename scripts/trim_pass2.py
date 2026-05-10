#!/usr/bin/env python3
"""Aggressive second-pass trim: 47 → ≤30 pages."""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os

BASE = r'E:\Desktop\高校机器人创意大赛'
DOC_PATH = os.path.join(BASE, '创意赛技术文档.docx')
doc = Document(DOC_PATH)


def delete_range(start, end):
    for i in range(end, start - 1, -1):
        doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
    return end - start + 1


def find_heading(text, level='Heading 1'):
    for i, para in enumerate(doc.paragraphs):
        if level in (para.style.name or '') and text in para.text:
            return i
    return None


def delete_image_and_caption(para_idx):
    """Delete an image paragraph and its following caption."""
    # para_idx is the image paragraph; caption is para_idx+1
    p_img = doc.paragraphs[para_idx]
    p_cap = doc.paragraphs[para_idx + 1]
    drawings = p_img._element.findall('.//' + qn('w:drawing'))
    if drawings and p_cap.text.strip().startswith('图'):
        p_cap._element.getparent().remove(p_cap._element)
        p_img._element.getparent().remove(p_img._element)
        return True
    return False


# ═══════════════════════════════════════════════════════════════
print('=== Pass 2: Aggressive trimming ===')

# 1. Remove 2 less-critical diagrams: vision pipeline (图4-8) and hand-eye calibration (图4-3)
#    These are already well-described in the text.
removed = 0
for i, para in enumerate(doc.paragraphs):
    pnext = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else None
    if pnext and ('图4-3' in pnext.text and '手眼标定' in pnext.text):
        if delete_image_and_caption(i):
            removed += 1
            print(f'  Removed 图4-3 (hand-eye calibration diagram)')
    if pnext and ('图4-8' in pnext.text and '视觉识别流水线' in pnext.text):
        if delete_image_and_caption(i):
            removed += 1
            print(f'  Removed 图4-8 (vision pipeline diagram)')

print(f'  Removed {removed} diagrams')


# 2. Aggressively trim Ch2: delete verbose paragraphs (>200 chars) that aren't essential
ch2 = find_heading('2. ', 'Heading 1')
ch3 = find_heading('3. ', 'Heading 1')
if ch2 and ch3:
    deleted = 0
    for i in range(ch3 - 1, ch2, -1):
        t = doc.paragraphs[i].text.strip()
        style = doc.paragraphs[i].style.name or ''
        if 'Heading' not in style and len(t) > 200 and deleted < 8:
            doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
            deleted += 1
    print(f'  Ch2: deleted {deleted} verbose paragraphs')


# 3. Trim Ch4: delete verbose paragraphs (>250 chars)
ch4 = find_heading('4. ', 'Heading 1')
ch5 = find_heading('5. ', 'Heading 1')
if ch4 and ch5:
    deleted = 0
    for i in range(ch5 - 1, ch4, -1):
        t = doc.paragraphs[i].text.strip()
        style = doc.paragraphs[i].style.name or ''
        if 'Heading' not in style and len(t) > 250 and deleted < 12:
            doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
            deleted += 1
    print(f'  Ch4: deleted {deleted} verbose paragraphs')


# 4. Remove 附录C installation diagrams (附图C-1, 附图C-2) - keep only experiment photos
for i, para in enumerate(doc.paragraphs):
    if ('附图C-1' in para.text and '侧视图' in para.text) or \
       ('附图C-2' in para.text and '俯视图' in para.text):
        prev = doc.paragraphs[i - 1]
        drawings = prev._element.findall('.//' + qn('w:drawing'))
        if drawings:
            para._element.getparent().remove(para._element)
            prev._element.getparent().remove(prev._element)
            print(f'  Removed {para.text.strip()[:60]}')


# 5. Compress references: keep first 10, delete rest
refs = find_heading('参考文献', 'Heading 1')
app = find_heading('附录', 'Heading 1')
if refs and app:
    ref_count = 0
    delete_start = None
    for i in range(refs + 1, app):
        t = doc.paragraphs[i].text.strip()
        if t.startswith('['):
            ref_count += 1
        if ref_count >= 10 and delete_start is None:
            delete_start = i + 1
    if delete_start and delete_start < app:
        n = delete_range(delete_start, app - 1)
        print(f'  References: kept 10, deleted {n} paragraphs')


# 6. Remove 附录C heading if section is now empty
app_c = find_heading('附录C', 'Heading 2')
app_d_idx = find_heading('附录D', 'Heading 2')  # deleted earlier, may not exist
if app_c:
    # Check if there's any content between C heading and next heading
    next_heading = None
    for i in range(app_c + 1, len(doc.paragraphs)):
        if 'Heading' in (doc.paragraphs[i].style.name or ''):
            next_heading = i
            break
    if next_heading and next_heading == app_c + 1:
        # Nothing between, delete heading too
        doc.paragraphs[app_c]._element.getparent().remove(doc.paragraphs[app_c]._element)
        print(f'  Removed empty 附录C heading')


# ═══════════════════════════════════════════════════════════════
doc.save(DOC_PATH)
print(f'\nSaved. Size: {os.path.getsize(DOC_PATH)/1024/1024:.1f} MB')
