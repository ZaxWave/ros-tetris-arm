#!/usr/bin/env python3
"""Fix cover blank page + trim long table descriptions."""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os, re

BASE = r'E:\Desktop\高校机器人创意大赛'
DOC_PATH = os.path.join(BASE, '创意赛技术文档.docx')
doc = Document(DOC_PATH)


# ═══════════════════════════════════════════════════════════════
# 1. Fix cover: remove excessive empty paragraphs (P0-P5 → keep 2)
#    Also fix order: title → subtitle → badge → team info
print('=== 1. Fix cover layout ===')

# Current: P0-P5 empty, P6 title, P7 empty, P8 badge, P9 subtitle
# Wanted:  2 empty → title → empty → subtitle → badge → 3 empty → team info
# Delete P0-P5 (first 6 empty paras)
for i in range(5, -1, -1):
    doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)


# Now P0=title, P1=empty, P2=badge, P3=subtitle
# Fix: swap P2 (badge) and P3 (subtitle)
# After fix: P0=title, P1=empty, P2=subtitle, P3=badge
# The badge was inserted at the right place originally but the subtitle shifted after insertions.

# Find subtitle paragraph ("专项赛一")
subtitle_idx = None
badge_idx = None
for i, p in enumerate(doc.paragraphs[:15]):
    if '专项赛一' in p.text:
        subtitle_idx = i
    if p._element.findall('.//'+qn('w:drawing')):
        badge_idx = i

if subtitle_idx and badge_idx:
    if badge_idx < subtitle_idx:
        # Badge is BEFORE subtitle, need to swap
        # Move badge element after subtitle
        badge_elem = doc.paragraphs[badge_idx]._element
        subtitle_elem = doc.paragraphs[subtitle_idx]._element
        parent = badge_elem.getparent()

        # Remove badge from current position
        parent.remove(badge_elem)
        # Insert badge after subtitle
        subtitle_pos = list(parent).index(subtitle_elem)
        parent.insert(subtitle_pos + 1, badge_elem)
        print(f'  Fixed: badge now after subtitle')

# Clean excessive empty paragraphs between cover elements
# Keep: title, 1 empty, subtitle, badge, 3 empty, team info
cover_range = min(20, len(doc.paragraphs))
empties_in_row = 0
for i in range(cover_range - 1, -1, -1):
    t = doc.paragraphs[i].text.strip()
    has_img = bool(doc.paragraphs[i]._element.findall('.//'+qn('w:drawing')))
    s = doc.paragraphs[i].style.name or ''
    if 'Heading' in s:
        break
    if not t and not has_img:
        empties_in_row += 1
        if empties_in_row > 3:
            doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
    else:
        empties_in_row = 0

print('  Cover spacing cleaned')

# ═══════════════════════════════════════════════════════════════
# 2. Trim long table cell descriptions
print('\n=== 2. Trim table text ===')
total_cells_trimmed = 0
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip()
            if len(text) > 100:
                # Shorten by removing verbose details
                # Common patterns to shorten
                shortened = text
                # Remove parenthetical details
                shortened = re.sub(r'（[^）]*）', '', shortened)
                shortened = re.sub(r'\([^)]*\)', '', shortened)
                # Truncate very long descriptions
                if len(shortened) > 120:
                    # Keep first sentence or first 100 chars
                    sentences = re.split(r'[。；;]', shortened)
                    if len(sentences) > 1 and len(sentences[0]) > 20:
                        shortened = sentences[0] + '。'
                    else:
                        shortened = shortened[:100] + '…'
                shortened = shortened.strip()
                if len(shortened) < len(text) - 5:
                    # Write shortened text back
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = ''
                        if para.runs:
                            para.runs[0].text = shortened
                        else:
                            para.text = shortened
                    total_cells_trimmed += 1

print(f'  Trimmed {total_cells_trimmed} table cells')

# ═══════════════════════════════════════════════════════════════
# 3. Remove extra text in Ch4
print('\n=== 3. More Ch4 trimming ===')
ch4 = ch5 = None
for i, p in enumerate(doc.paragraphs):
    s = p.style.name or ''
    if 'Heading 1' in s and '4. ' in p.text: ch4 = i
    if 'Heading 1' in s and '5. ' in p.text: ch5 = i

if ch4 and ch5:
    deleted = 0
    for i in range(ch5 - 1, ch4, -1):
        t = doc.paragraphs[i].text.strip()
        s = doc.paragraphs[i].style.name or ''
        if 'Heading' in s: continue
        if len(t) > 90:
            pnext = doc.paragraphs[i+1] if i+1 < len(doc.paragraphs) else None
            if pnext and pnext._element.findall('.//'+qn('w:drawing')): continue
            doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
            deleted += 1
            if deleted >= 8: break
    print(f'  Deleted {deleted} paragraphs from Ch4')

# ═══════════════════════════════════════════════════════════════
doc.save(DOC_PATH)
print(f'\nSaved. Size: {os.path.getsize(DOC_PATH)/1024/1024:.1f} MB')
