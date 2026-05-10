#!/usr/bin/env python3
"""
Complete document rebuild:
1. Remove old page breaks, clean up empty paragraphs
2. Add school badge to cover, fix cover spacing
3. Add page breaks at correct locations (no blank pages)
4. Insert TikZ diagrams (academic LaTeX quality)
5. Insert experiment photos
6. Trim appendix code
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import os

BASE = r'E:\Desktop\高校机器人创意大赛'
DOC_PATH = os.path.join(BASE, '创意赛技术文档.docx')
DIAGRAMS = os.path.join(BASE, 'assets', 'diagrams')
PHOTOS = os.path.join(BASE, 'assets', 'images')
BADGE = os.path.join(PHOTOS, 'school_badge.png')

WML_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

doc = Document(DOC_PATH)


def remove_page_breaks():
    """Remove all page breaks from runs and paragraph properties."""
    count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            for br in list(run._element.findall(qn('w:br'))):
                if br.get(qn('w:type')) == 'page':
                    run._element.remove(br)
                    count += 1
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            for br in list(pPr.findall(qn('w:br'))):
                if br.get(qn('w:type')) == 'page':
                    pPr.remove(br)
                    count += 1
    print(f'  Removed {count} page breaks')
    return count


def add_page_break(para):
    """Add page break to paragraph properties."""
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr xmlns:w="{WML_NS}"></w:pPr>')
        para._element.insert(0, pPr)
    for br in pPr.findall(qn('w:br')):
        if br.get(qn('w:type')) == 'page':
            return
    br = parse_xml(f'<w:br xmlns:w="{WML_NS}" w:type="page"/>')
    pPr.append(br)


def delete_empty_paras_before(idx, max_lookback=6):
    """Delete empty paragraphs before a given paragraph, keeping at most 1."""
    deleted = 0
    for j in range(idx - 1, max(idx - max_lookback - 1, -1), -1):
        t = doc.paragraphs[j].text.strip()
        has_img = bool(doc.paragraphs[j]._element.findall('.//' + qn('w:drawing')))
        if not t and not has_img:
            if deleted >= 1:  # already found one empty, delete rest
                doc.paragraphs[j]._element.getparent().remove(doc.paragraphs[j]._element)
                deleted += 1
            else:
                deleted += 1  # keep first empty
        else:
            break
    return deleted


def insert_centered_image_after(para_index, image_path, width=Inches(1.3)):
    """Insert centered image after given paragraph."""
    para = doc.paragraphs[para_index]
    parent = para._element.getparent()
    insert_idx = list(parent).index(para._element) + 1

    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.paragraph_format.space_before = Pt(8)
    img_p.paragraph_format.space_after = Pt(8)
    run = img_p.add_run()
    run.add_picture(image_path, width=width)

    img_p._element.getparent().remove(img_p._element)
    parent.insert(insert_idx, img_p._element)


def insert_image_with_caption(para_index, image_path, label, desc, width=Inches(5.5)):
    """Insert figure + caption after paragraph."""
    para = doc.paragraphs[para_index]
    parent = para._element.getparent()
    idx = list(parent).index(para._element) + 1

    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.paragraph_format.space_before = Pt(8)
    img_p.paragraph_format.space_after = Pt(2)
    img_p.add_run().add_picture(image_path, width=width)

    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_after = Pt(10)
    lbl = cap_p.add_run(label + '  ')
    lbl.bold = True; lbl.font.size = Pt(9); lbl.font.name = 'Times New Roman'
    dsc = cap_p.add_run(desc)
    dsc.font.size = Pt(9); dsc.font.name = 'Times New Roman'

    cap_p._element.getparent().remove(cap_p._element)
    img_p._element.getparent().remove(img_p._element)
    parent.insert(idx, img_p._element)
    parent.insert(idx + 1, cap_p._element)
    print(f'  {label} after P{para_index}')


# ═══════════════════════════════════════════════════════════════
print('=== 1. Remove old page breaks ===')
remove_page_breaks()

# ═══════════════════════════════════════════════════════════════
print('\n=== 2. Fix cover — school badge + spacing ===')
# P5: title, P7: subtitle, insert badge after P7
insert_centered_image_after(7, BADGE, width=Inches(1.3))
# Clean excessive blank lines on cover (now P8 has badge, reduce spacing)
# Keep P5 (title), P6 (empty), P7 (subtitle), P8 (badge), then 4 blank → team info

# ═══════════════════════════════════════════════════════════════
print('\n=== 3. Clean empty paragraphs near headings ===')
# Find all Heading 1 paragraphs and clean empties before them
heading_indices = []
for i, para in enumerate(doc.paragraphs):
    if 'Heading 1' in (para.style.name or ''):
        heading_indices.append(i)
        delete_empty_paras_before(i, max_lookback=4)

print(f'  Cleaned around {len(heading_indices)} chapter headings')

# ═══════════════════════════════════════════════════════════════
print('\n=== 4. Add strategic page breaks ===')
# Find key locations
breaks_at = {}
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    s = para.style.name or ''
    if t == '目录' and 'Heading' in s: breaks_at['TOC'] = i
    elif t == '摘要' and 'Heading' in s: breaks_at['Abstract'] = i
    elif t.startswith('1. ') and 'Heading 1' in s: breaks_at['Ch1'] = i
    elif t.startswith('2. ') and 'Heading 1' in s: breaks_at['Ch2'] = i
    elif t.startswith('3. ') and 'Heading 1' in s: breaks_at['Ch3'] = i
    elif t.startswith('4. ') and 'Heading 1' in s: breaks_at['Ch4'] = i
    elif t.startswith('5. ') and 'Heading 1' in s: breaks_at['Ch5'] = i
    elif t.startswith('6. ') and 'Heading 1' in s: breaks_at['Ch6'] = i
    elif t == '参考文献' and 'Heading' in s: breaks_at['Refs'] = i
    elif t == '附录' and 'Heading' in s: breaks_at['App'] = i

for name, idx in breaks_at.items():
    add_page_break(doc.paragraphs[idx])
    print(f'  Page break → {name} (P{idx})')

# ═══════════════════════════════════════════════════════════════
print('\n=== 5-6. Insert all images (single pass, bottom-to-top) ===')

all_insertions = [
    # Diagrams
    (297, 'diagrams', 'fig5_pick_place_sequence.png', '图4-6', '拾取-放置运动时序图（9.8 s单次循环）', Inches(5.8)),
    (263, 'diagrams', 'fig4_decision_flow.png', '图4-5', '混合决策策略流程（Dellacherie启发式筛选 + DQN长期价值优化）', Inches(5.8)),
    (251, 'diagrams', 'fig6_dqn_architecture.png', '图4-4', '深度Q网络架构（211维状态 $\rightarrow$ 3层MLP $\rightarrow$ 400维Q值）', Inches(5.8)),
    (218, 'diagrams', 'fig3_vision_pipeline.png', '图4-8', '视觉识别流水线（D435采集$\rightarrow$YOLOv8分割$\rightarrow$模板匹配$\rightarrow$6-DOF位姿）', Inches(5.8)),
    (193, 'diagrams', 'fig7_handeye_calibration.png', '图4-3', '手眼标定流程（Tsai-Lenz算法 $AX=XB$ 求解）', Inches(5.3)),
    (134, 'diagrams', 'fig2_ros_node_graph.png', '图2-1', 'ROS节点通信架构图（核心节点与话题/服务关系）', Inches(5.5)),
    (102, 'diagrams', 'fig1_system_architecture.png', '图1-2', '系统总体架构图（感知$\rightarrow$决策$\rightarrow$执行三层递阶控制）', Inches(5.8)),
    # Experiment photos
    (325, 'photos', 'assembly_result.jpg', '图5-4', '35块俄罗斯方块全自动拼接完成效果', Inches(5.0)),
    (317, 'photos', 'grasping_process.jpg', '图5-3', '气动吸盘精准拾取俄罗斯方块过程', Inches(5.0)),
    (311, 'photos', 'recognition_diagram.jpg', '图5-2', '俄罗斯方块视觉识别效果（YOLOv8实例分割+分类标注）', Inches(5.0)),
    (308, 'photos', 'system_running.jpg', '图5-1', '系统整体运行场景（RM65-B机械臂 + D435相机 + 工作台）', Inches(5.0)),
]

# Sort by para_index descending (bottom-to-top so indices stay valid)
all_insertions.sort(key=lambda x: x[0], reverse=True)
for idx, src, fname, label, desc, w in all_insertions:
    src_dir = DIAGRAMS if src == 'diagrams' else PHOTOS
    insert_image_with_caption(idx, os.path.join(src_dir, fname), label, desc, w)

# ═══════════════════════════════════════════════════════════════
print('\n=== 7. Trim appendix code ===')
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    s = para.style.name or ''
    if 'Heading 3' in s and ('A.1' in t or 'A.2' in t or 'A.3' in t):
        lines = 0
        for j in range(i + 1, min(i + 100, len(doc.paragraphs))):
            pj = doc.paragraphs[j]
            if 'Heading' in (pj.style.name or ''):
                break
            if pj.text.strip():
                lines += 1
                if lines > 25:
                    # Delete from here to next heading
                    end = j
                    for k in range(j, min(j + 200, len(doc.paragraphs))):
                        if 'Heading' in (doc.paragraphs[k].style.name or ''):
                            break
                        end = k
                    for d in range(end, j - 1, -1):
                        doc.paragraphs[d]._element.getparent().remove(doc.paragraphs[d]._element)
                    print(f'  Trimmed {end - j + 1} lines from {t[:50]}')
                    break

# ═══════════════════════════════════════════════════════════════
print('\n=== 8. Save ===')
doc.save(DOC_PATH)
size_mb = os.path.getsize(DOC_PATH) / 1024 / 1024
print(f'  Saved: {DOC_PATH} ({size_mb:.1f} MB)')
print('Done!')
