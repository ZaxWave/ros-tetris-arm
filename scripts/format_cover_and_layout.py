#!/usr/bin/env python3
"""
Format document: add school badge to cover, insert strategic page breaks,
trim verbose content for conciseness.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import os

BASE = r'E:\Desktop\高校机器人创意大赛'
DOC_PATH = os.path.join(BASE, '创意赛技术文档.docx')
BADGE_PATH = os.path.join(BASE, 'assets', 'images', 'school_badge.png')

doc = Document(DOC_PATH)


WML_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def add_page_break_before(para_index):
    """Add page break before the given paragraph."""
    para = doc.paragraphs[para_index]
    pPr = para._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr xmlns:w="{WML_NS}"></w:pPr>')
        para._element.insert(0, pPr)
    # Check if page break already exists
    existing = pPr.findall(qn('w:br'))
    for br in existing:
        if br.get(qn('w:type')) == 'page':
            return  # already has page break
    br = parse_xml(f'<w:br xmlns:w="{WML_NS}" w:type="page"/>')
    pPr.append(br)
    print(f'  Page break before P{para_index}: {para.text[:60]}')


def insert_centered_image(para_index, image_path, width=Inches(1.2)):
    """Insert a centered image after the given paragraph."""
    para = doc.paragraphs[para_index]
    parent = para._element.getparent()
    insert_idx = list(parent).index(para._element) + 1

    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.paragraph_format.space_before = Pt(6)
    img_p.paragraph_format.space_after = Pt(6)
    run = img_p.add_run()
    run.add_picture(image_path, width=width)

    # Move to correct position
    img_p._element.getparent().remove(img_p._element)
    parent.insert(insert_idx, img_p._element)
    print(f'  School badge inserted after P{para_index}')
    return img_p


def delete_paragraphs(start, end):
    """Delete paragraphs from start to end (inclusive)."""
    for i in range(end, start - 1, -1):
        p = doc.paragraphs[i]
        p._element.getparent().remove(p._element)
    print(f'  Deleted P{start}-P{end}')


# ═══════════════════════════════════════════════════════════════
# STEP 1: Insert school badge on cover (after P7 subtitle, before P14 team)
# ═══════════════════════════════════════════════════════════════
print('=== Cover Page ===')
# P7 = subtitle line "专项赛一 · 基于ROS的单臂机器人\n俄罗斯方块自动化拼接系统"
insert_centered_image(7, BADGE_PATH, width=Inches(1.3))

# Clean up excessive empty paras on cover (P8-P13 are empty, reduce to P8-P10)
# After inserting badge at P8, P9-P14 shift to P9-P15. Keep 2 blank lines.


# ═══════════════════════════════════════════════════════════════
# STEP 2: Add page breaks at key transitions
# ═══════════════════════════════════════════════════════════════
print('\n=== Page Breaks ===')

# Find current paragraph indices for key headings (they may have shifted from image insertions)
toc_idx = None
abstract_idx = None
ch1_idx = None
ch2_idx = None
ch3_idx = None
ch4_idx = None
ch5_idx = None
ch6_idx = None
ref_idx = None
app_idx = None

for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    style = para.style.name
    if t == '目录' and 'Heading' in style:
        toc_idx = i
    elif t == '摘要' and 'Heading' in style:
        abstract_idx = i
    elif t.startswith('1. ') and 'Heading 1' in style:
        ch1_idx = i
    elif t.startswith('2. ') and 'Heading 1' in style:
        ch2_idx = i
    elif t.startswith('3. ') and 'Heading 1' in style:
        ch3_idx = i
    elif t.startswith('4. ') and 'Heading 1' in style:
        ch4_idx = i
    elif t.startswith('5. ') and 'Heading 1' in style:
        ch5_idx = i
    elif t.startswith('6. ') and 'Heading 1' in style:
        ch6_idx = i
    elif t == '参考文献' and 'Heading 1' in style:
        ref_idx = i
    elif t == '附录' and 'Heading 1' in style:
        app_idx = i

print(f'  TOC: P{toc_idx}, Abstract: P{abstract_idx}')
print(f'  Ch1: P{ch1_idx}, Ch2: P{ch2_idx}, Ch3: P{ch3_idx}, Ch4: P{ch4_idx}, Ch5: P{ch5_idx}, Ch6: P{ch6_idx}')
print(f'  References: P{ref_idx}, Appendix: P{app_idx}')

# Page break: after cover → before TOC
if toc_idx:
    add_page_break_before(toc_idx)

# Page break: after TOC → before Abstract
if abstract_idx:
    add_page_break_before(abstract_idx)

# Page break: after Abstract → before Chapter 1
if ch1_idx:
    add_page_break_before(ch1_idx)

# Page break between chapters
for idx in [ch2_idx, ch3_idx, ch4_idx, ch5_idx, ch6_idx]:
    if idx:
        add_page_break_before(idx)

# Page break before references
if ref_idx:
    add_page_break_before(ref_idx)

# Page break before appendix
if app_idx:
    add_page_break_before(app_idx)


# ═══════════════════════════════════════════════════════════════
# STEP 3: Trim verbose content
# ═══════════════════════════════════════════════════════════════
print('\n=== Trimming Content ===')

# Strategy: identify overly long paragraphs and trim redundancies
# The appendix code sections (A.1, A.2, A.3) are very long code listings.
# We'll keep the first ~30 lines of each code block as examples.

# Find appendix code sections
code_sections = []
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    style = para.style.name
    if 'Heading 3' in style and ('A.1' in t or 'A.2' in t or 'A.3' in t):
        code_sections.append((i, t))

print(f'  Found {len(code_sections)} code sections')

# For each code section, trim after the heading + about 25 lines of code
# Code in appendix is in Normal style paragraphs after the heading
for sec_idx, sec_title in code_sections:
    lines_kept = 0
    for j in range(sec_idx + 1, min(sec_idx + 100, len(doc.paragraphs))):
        p = doc.paragraphs[j]
        t = p.text.strip()
        style = p.style.name
        # Stop at next heading or after 25 code lines
        if 'Heading' in style:
            break
        if t == '':
            continue
        lines_kept += 1
        if lines_kept > 25:
            # Delete remaining code lines until next heading
            delete_start = j
            delete_end = j
            for k in range(j, min(j + 200, len(doc.paragraphs))):
                pk = doc.paragraphs[k]
                if 'Heading' in (pk.style.name or ''):
                    break
                delete_end = k
            if delete_end > delete_start:
                delete_paragraphs(delete_start, delete_end)
                print(f'  Trimmed {delete_end - delete_start + 1} lines from {sec_title}')
            break


# ═══════════════════════════════════════════════════════════════
# STEP 4: Save
# ═══════════════════════════════════════════════════════════════
doc.save(DOC_PATH)
import os
size = os.path.getsize(DOC_PATH) / 1024 / 1024
print(f'\nSaved: {DOC_PATH} ({size:.1f} MB)')
print('Done!')
