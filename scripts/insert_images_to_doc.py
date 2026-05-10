#!/usr/bin/env python3
"""
Insert professional diagrams and experiment photos into the competition document.
Processes insertions from bottom-to-top to preserve paragraph indices.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

BASE = r'E:\Desktop\高校机器人创意大赛'
DOC_PATH = os.path.join(BASE, '创意赛技术文档.docx')
OUT_PATH = os.path.join(BASE, '创意赛技术文档.docx')

DIAGRAMS = os.path.join(BASE, 'assets', 'diagrams')
PHOTOS = os.path.join(BASE, 'assets', 'images')

doc = Document(DOC_PATH)


def insert_at(doc, para_index, image_path, label, desc, width=Inches(5.5)):
    """Insert image + caption at a specific paragraph index.
    Works correctly when called bottom-to-top (descending para_index)."""
    body = doc.element.body
    para_elem = doc.paragraphs[para_index]._element

    # ── Build image paragraph ──
    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.paragraph_format.space_before = Pt(8)
    img_p.paragraph_format.space_after = Pt(2)

    run = img_p.add_run()
    run.add_picture(image_path, width=width)

    # ── Build caption paragraph ──
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_after = Pt(10)

    lbl = cap_p.add_run(label + '  ')
    lbl.bold = True
    lbl.font.size = Pt(9)
    lbl.font.name = 'Times New Roman'

    dsc = cap_p.add_run(desc)
    dsc.font.size = Pt(9)
    dsc.font.name = 'Times New Roman'

    # Remove from end
    cap_p._element.getparent().remove(cap_p._element)
    img_p._element.getparent().remove(img_p._element)

    # Insert after target
    target = para_elem
    target.getparent().insert(list(target.getparent()).index(target) + 1, img_p._element)
    target.getparent().insert(list(target.getparent()).index(target) + 2, cap_p._element)

    print(f'  {label} after P{para_index}')


# ═══════════════════════════════════════════════════════════════
# INSERTIONS (bottom → top so indices stay stable)
# ═══════════════════════════════════════════════════════════════

print('Inserting images bottom-to-top...')

# ── Experiment Photos (Section 5, bottom of doc first) ──

insert_at(doc, 325,
    os.path.join(PHOTOS, 'assembly_result.jpg'),
    '图5-4', '35块俄罗斯方块全自动拼接完成效果',
    width=Inches(5.0))

insert_at(doc, 317,
    os.path.join(PHOTOS, 'grasping_process.jpg'),
    '图5-3', '气动吸盘精准拾取俄罗斯方块过程',
    width=Inches(5.0))

insert_at(doc, 311,
    os.path.join(PHOTOS, 'recognition_diagram.jpg'),
    '图5-2', '俄罗斯方块视觉识别效果（YOLOv8实例分割+分类标注）',
    width=Inches(5.0))

insert_at(doc, 308,
    os.path.join(PHOTOS, 'system_running.jpg'),
    '图5-1', '系统整体运行场景（RM65-B机械臂 + D435相机 + 工作台）',
    width=Inches(5.0))

# ── Diagrams (Section 4) ──

insert_at(doc, 297,
    os.path.join(DIAGRAMS, 'fig5_pick_place_sequence.png'),
    '图4-5', '拾取-放置运动时序图（9.8s单次循环）',
    width=Inches(5.8))

insert_at(doc, 263,
    os.path.join(DIAGRAMS, 'fig4_decision_flow.png'),
    '图4-4', '混合决策策略流程（Dellacherie启发式筛选 + DQN长期价值优化）',
    width=Inches(5.8))

insert_at(doc, 251,
    os.path.join(DIAGRAMS, 'fig6_dqn_architecture.png'),
    '图4-3', 'DQN深度Q网络架构（211维状态 → 3层MLP → 400维动作价值）',
    width=Inches(5.8))

# Note: existing fig4-1/4-2 are at P226-230. We insert our pipeline BEFORE them as 图4-2
# and they become 图4-3/4-4. But our DQN is already 图4-3...
# Let me insert vision pipeline as a separate figure numbering.
# Actually let me insert vision pipeline after 4.2.1 heading (P207) as a different approach.

insert_at(doc, 218,
    os.path.join(DIAGRAMS, 'fig3_vision_pipeline.png'),
    '图4-8', '视觉识别流水线（D435采集→YOLOv8分割→模板匹配→手眼变换→抓取位姿）',
    width=Inches(5.8))

insert_at(doc, 193,
    os.path.join(DIAGRAMS, 'fig7_handeye_calibration.png'),
    '图4-3', '手眼标定流程（Tsai-Lenz算法 AX=XB 求解）',
    width=Inches(5.3))

# ── Early Diagrams ──

insert_at(doc, 134,
    os.path.join(DIAGRAMS, 'fig2_ros_node_graph.png'),
    '图2-1', 'ROS节点通信架构图（核心节点与话题/服务关系）',
    width=Inches(5.5))

insert_at(doc, 102,
    os.path.join(DIAGRAMS, 'fig1_system_architecture.png'),
    '图1-2', '系统总体架构图（感知→决策→执行三层递阶架构）',
    width=Inches(5.8))


# ── Save ──
doc.save(OUT_PATH)
print(f'\nSaved: {OUT_PATH}')
print('Done!')
