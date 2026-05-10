#!/usr/bin/env python3
"""Radical trim: delete Ch6, 附录C, more images, more text."""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os

BASE = r'E:\Desktop\高校机器人创意大赛'
DOC_PATH = os.path.join(BASE, '创意赛技术文档.docx')
doc = Document(DOC_PATH)

def delete_range(start, end):
    for i in range(end, start - 1, -1):
        doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
    return end - start + 1

def find_heading(text, level='Heading 1'):
    for i, p in enumerate(doc.paragraphs):
        if level in (p.style.name or '') and text in p.text:
            return i
    return None

def delete_image_and_caption(para_idx):
    if para_idx + 1 >= len(doc.paragraphs): return False
    pi = doc.paragraphs[para_idx]; pc = doc.paragraphs[para_idx + 1]
    if pi._element.findall('.//'+qn('w:drawing')) and pc.text.strip().startswith('图'):
        pc._element.getparent().remove(pc._element)
        pi._element.getparent().remove(pi._element)
        return True
    return False

# ═══════════════════════════════════════════════════════════════
# 1. Delete Chapter 6 entirely (优化与展望)
print('=== 1. Delete Chapter 6 ===')
ch6 = find_heading('6. ', 'Heading 1')
refs = find_heading('参考文献', 'Heading 1')
if ch6 and refs:
    n = delete_range(ch6, refs - 1)
    print(f'  Deleted Ch6: {n} paragraphs')

# ═══════════════════════════════════════════════════════════════
# 2. Delete 附录C entirely
print('\n=== 2. Delete Appendix C ===')
app_c = find_heading('附录C', 'Heading 2')
if app_c:
    # Find next heading or end
    end = len(doc.paragraphs) - 1
    for i in range(app_c + 1, len(doc.paragraphs)):
        if 'Heading' in (doc.paragraphs[i].style.name or ''):
            end = i - 1
            break
    n = delete_range(app_c, end)
    print(f'  Deleted Appendix C: {n} paragraphs')

# ═══════════════════════════════════════════════════════════════
# 3. Remove 图5-2 (redundant with 图4-1 YOLO detection)
print('\n=== 3. Remove redundant images ===')
for i, p in enumerate(doc.paragraphs):
    pnext = doc.paragraphs[i+1] if i+1 < len(doc.paragraphs) else None
    if pnext and '图5-2' in pnext.text and '视觉识别效果' in pnext.text:
        if delete_image_and_caption(i): print('  Removed 图5-2 (duplicate with 图4-1)')

# ═══════════════════════════════════════════════════════════════
# 4. Aggressive Ch4 trim: delete ALL paragraphs >120 chars
print('\n=== 4. Aggressive Ch4 text trim ===')
ch4 = find_heading('4. ', 'Heading 1')
ch5 = find_heading('5. ', 'Heading 1')
if ch4 and ch5:
    deleted = 0
    for i in range(ch5 - 1, ch4, -1):
        t = doc.paragraphs[i].text.strip()
        s = doc.paragraphs[i].style.name or ''
        if 'Heading' in s or len(t) <= 120:
            continue
        if t.startswith('图') or t.startswith('表'):
            continue
        # Don't delete if followed by image
        pnext = doc.paragraphs[i+1] if i+1 < len(doc.paragraphs) else None
        if pnext and pnext._element.findall('.//'+qn('w:drawing')):
            continue
        doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
        deleted += 1
    print(f'  Ch4: deleted {deleted} paragraphs (120+ chars)')

# ═══════════════════════════════════════════════════════════════
# 5. Trim Ch2 more
print('\n=== 5. Trim Ch2 more ===')
ch2 = find_heading('2. ', 'Heading 1')
ch3 = find_heading('3. ', 'Heading 1')
if ch2 and ch3:
    deleted = 0
    for i in range(ch3 - 1, ch2, -1):
        t = doc.paragraphs[i].text.strip()
        s = doc.paragraphs[i].style.name or ''
        if 'Heading' in s or len(t) <= 100:
            continue
        if t.startswith('图') or t.startswith('表'):
            continue
        doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
        deleted += 1
    print(f'  Ch2: deleted {deleted} paragraphs (100+ chars)')

# ═══════════════════════════════════════════════════════════════
doc.save(DOC_PATH)
print(f'\nSaved. Size: {os.path.getsize(DOC_PATH)/1024/1024:.1f} MB')
