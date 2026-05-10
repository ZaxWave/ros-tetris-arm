#!/usr/bin/env python3
"""Aggressive trim to hit ≤30 pages. Remove diagrams, verbose text, compress refs."""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os

BASE = r'E:\Desktop\高校机器人创意大赛'
DOC_PATH = os.path.join(BASE, '创意赛技术文档.docx')
doc = Document(DOC_PATH)

def delete_range(start, end):
    for i in range(end, start - 1, -1):
        doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)

def find_heading(text, level='Heading 1'):
    for i, p in enumerate(doc.paragraphs):
        if level in (p.style.name or '') and text in p.text:
            return i
    return None

def delete_image_and_caption(para_idx):
    """Delete image at para_idx and its caption at para_idx+1."""
    if para_idx + 1 >= len(doc.paragraphs):
        return False
    p_img = doc.paragraphs[para_idx]
    p_cap = doc.paragraphs[para_idx + 1]
    if p_img._element.findall('.//'+qn('w:drawing')) and p_cap.text.strip().startswith('图'):
        p_cap._element.getparent().remove(p_cap._element)
        p_img._element.getparent().remove(p_img._element)
        return True
    return False

removed_total = 0

# ═══════════════════════════════════════════════════════════════
# 1. Remove 图4-4 (DQN) and 图4-5 (Decision flow) — text describes them
print('=== 1. Remove 2 redundant diagrams ===')
for i, p in enumerate(doc.paragraphs):
    pnext = doc.paragraphs[i+1] if i+1 < len(doc.paragraphs) else None
    if pnext:
        if ('图4-4' in pnext.text and 'DQN' in pnext.text) or \
           ('图4-4' in pnext.text and '深度Q' in pnext.text):
            if delete_image_and_caption(i):
                print('  Removed 图4-4 (DQN architecture)')
        if ('图4-5' in pnext.text and '混合决策' in pnext.text):
            if delete_image_and_caption(i):
                print('  Removed 图4-5 (decision flow)')

# 2. Remove 附图C-2
for i, p in enumerate(doc.paragraphs):
    if '附图C-2' in p.text and '俯视图' in p.text:
        prev = doc.paragraphs[i-1]
        if prev._element.findall('.//'+qn('w:drawing')):
            p._element.getparent().remove(p._element)
            prev._element.getparent().remove(prev._element)
            print('  Removed 附图C-2')

# ═══════════════════════════════════════════════════════════════
# 3. Aggressive Ch4 trim: delete long paragraphs (>180 chars)
print('\n=== 3. Trim Ch4 aggressively ===')
ch4 = find_heading('4. ', 'Heading 1')
ch5 = find_heading('5. ', 'Heading 1')
if ch4 and ch5:
    deleted = 0
    for i in range(ch5 - 1, ch4, -1):
        t = doc.paragraphs[i].text.strip()
        s = doc.paragraphs[i].style.name or ''
        # Skip headings, short paras, image captions, tables
        if 'Heading' in s or len(t) < 180:
            continue
        if t.startswith('图') or t.startswith('表'):
            continue
        # Check if next para is an image (don't delete the text before an image)
        pnext = doc.paragraphs[i+1] if i+1 < len(doc.paragraphs) else None
        if pnext and pnext._element.findall('.//'+qn('w:drawing')):
            continue
        doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
        deleted += 1
        if deleted >= 15:
            break
    print(f'  Deleted {deleted} verbose paragraphs from Ch4')

# ═══════════════════════════════════════════════════════════════
# 4. Trim Ch2 verbose paragraphs
print('\n=== 4. Trim Ch2 ===')
ch2 = find_heading('2. ', 'Heading 1')
ch3 = find_heading('3. ', 'Heading 1')
if ch2 and ch3:
    deleted = 0
    for i in range(ch3 - 1, ch2, -1):
        t = doc.paragraphs[i].text.strip()
        s = doc.paragraphs[i].style.name or ''
        if 'Heading' in s or len(t) < 150:
            continue
        if t.startswith('图') or t.startswith('表'):
            continue
        doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
        deleted += 1
        if deleted >= 6:
            break
    print(f'  Deleted {deleted} verbose paragraphs from Ch2')

# ═══════════════════════════════════════════════════════════════
# 5. Compress references: keep only first 8
print('\n=== 5. Compress references ===')
refs = find_heading('参考文献', 'Heading 1')
app = find_heading('附录', 'Heading 1')
if refs and app:
    rcount = 0; cut_from = None
    for i in range(refs + 1, app):
        t = doc.paragraphs[i].text.strip()
        if t.startswith('['):
            rcount += 1
        if rcount >= 8 and cut_from is None:
            cut_from = i + 1
    if cut_from and cut_from < app:
        n = app - cut_from
        delete_range(cut_from, app - 1)
        print(f'  References: kept 8, deleted {n} paragraphs')

# ═══════════════════════════════════════════════════════════════
# 6. Remove 附录C if empty
for i, p in enumerate(doc.paragraphs):
    if '附录C' in p.text and 'Heading 2' in (p.style.name or ''):
        # Check next paragraph
        nxt = doc.paragraphs[i+1] if i+1 < len(doc.paragraphs) else None
        if nxt and 'Heading' in (nxt.style.name or ''):
            p._element.getparent().remove(p._element)
            print('  Removed empty 附录C heading')

# ═══════════════════════════════════════════════════════════════
doc.save(DOC_PATH)
print(f'\nSaved. Size: {os.path.getsize(DOC_PATH)/1024/1024:.1f} MB')
