# -*- coding: utf-8 -*-
"""生成�?��届高校智能机器人创意大赛�?�?���?v3.0
- 三线表格�?- 保留原文档图�?- 全面升级�?�?���?"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from latex2mathml import converter as latex2mathml_converter
import mathml2omml
from lxml import etree
import os

NS_M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

IMG_DIR = r'E:\Desktop\高校机器人创意大赛\images'

doc = Document()

# 设置 Word 打开时自动更新域（TOC �?��生成�?settings = doc.settings.element
update_fields = OxmlElement('w:updateFields')
update_fields.set(qn('w:val'), 'true')
settings.append(update_fields)

# �?�? 编号计数�?�?�?
_chapter_num = 1
_table_counter = {}
_figure_counter = {}
_appendix_letter = None
_appendix_table_counter = 0
_appendix_figure_counter = 0

def set_chapter(num):
    global _chapter_num, _appendix_letter
    _chapter_num = num
    _appendix_letter = None

def set_appendix(letter):
    global _appendix_letter, _appendix_table_counter, _appendix_figure_counter
    _appendix_letter = letter
    _appendix_table_counter = 0
    _appendix_figure_counter = 0

def _next_table_num():
    global _appendix_table_counter
    if _appendix_letter:
        _appendix_table_counter += 1
        return f'附表{_appendix_letter}-{_appendix_table_counter}'
    key = _chapter_num
    _table_counter[key] = _table_counter.get(key, 0) + 1
    return f'表{key}-{_table_counter[key]}'

def _next_figure_num():
    global _appendix_figure_counter
    if _appendix_letter:
        _appendix_figure_counter += 1
        return f'附图{_appendix_letter}-{_appendix_figure_counter}'
    key = _chapter_num
    _figure_counter[key] = _figure_counter.get(key, 0) + 1
    return f'图{key}-{_figure_counter[key]}'

# �?�? 全局样式 �?�?
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# �?�? 标�?不使用内�?��式，手动格式化确保字体�?�?�?�?
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sz = {1: Pt(16), 2: Pt(14), 3: Pt(13)}.get(level, Pt(14))
    sp_b = {1: Pt(18), 2: Pt(12), 3: Pt(8)}.get(level, Pt(12))
    sp_a = {1: Pt(12), 2: Pt(8), 3: Pt(6)}.get(level, Pt(6))
    h.paragraph_format.space_before = sp_b
    h.paragraph_format.space_after = sp_a
    h.paragraph_format.line_spacing = 1.25
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = sz
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def add_para(text, bold=False, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    run.bold = bold
    return p


def add_formula(latex_str, display=True):
    """Add a LaTeX equation as native OMML equation to the document."""
    try:
        mathml = latex2mathml_converter.convert(latex_str)
        omml_str = mathml2omml.convert(mathml)
        if display:
            wrapped = f'<m:oMathPara xmlns:m="{NS_M}"><m:oMathParaPr><m:jc m:val="center"/></m:oMathParaPr>{omml_str}</m:oMathPara>'
        else:
            wrapped = f'<m:oMath xmlns:m="{NS_M}">{omml_str}</m:oMath>'
        elem = etree.fromstring(wrapped)
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if display else WD_ALIGN_PARAGRAPH.LEFT
        para._p.append(elem)
        return para
    except Exception:
        # Fallback to plain text if OMML conversion fails
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(latex_str)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.italic = True
        return p


def add_toc():
    """Add a Table of Contents field (Heading 1-3). Auto-populates on F9 in Word."""
    para = doc.add_paragraph()
    run = para.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_char_begin)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = r' TOC \o "1-3" \h \z \u '
    run._r.append(instr)
    fld_char_separate = OxmlElement('w:fldChar')
    fld_char_separate.set(qn('w:fldCharType'), 'separate')
    run._r.append(fld_char_separate)
    run2 = para.add_run('（�?在Word�?���?��击�?处，选择"更新�?以生成目录）')
    run2.font.name = '宋体'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run2.font.size = Pt(12)
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')
    run2._r.append(fld_char_end)
    return para


def add_code_block(code_text):
    """Add a code block with gray background, one paragraph per line for proper alignment."""
    lines = code_text.split('\n')
    paras = []
    for idx, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.right_indent = Cm(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        # 灰色底纹
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F5F5F5')
        p._p.get_or_add_pPr().append(shd)
        # 段落边�?（仅首�?加上边�?，末行加下边框）
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        # 左�?�右边�?
        for side in ['left', 'right']:
            s = OxmlElement(f'w:{side}')
            s.set(qn('w:val'), 'single')
            s.set(qn('w:sz'), '4')
            s.set(qn('w:space'), '4')
            s.set(qn('w:color'), 'C0C0C0')
            pbdr.append(s)
        if idx == 0:
            # 首�?加上边�?
            s = OxmlElement('w:top')
            s.set(qn('w:val'), 'single')
            s.set(qn('w:sz'), '4')
            s.set(qn('w:space'), '4')
            s.set(qn('w:color'), 'C0C0C0')
            pbdr.append(s)
        if idx == len(lines) - 1:
            # �??加下边�?
            s = OxmlElement('w:bottom')
            s.set(qn('w:val'), 'single')
            s.set(qn('w:sz'), '4')
            s.set(qn('w:space'), '4')
            s.set(qn('w:color'), 'C0C0C0')
            pbdr.append(s)
        pPr.append(pbdr)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        paras.append(p)
    return paras


def set_three_line_table(table):
    """将表格�?�?��三线表格式：顶线粗�?�表头下线细、底线粗，无竖线无内�?��"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

    # 移除�?有先前的边�?设置
    for existing_borders in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing_borders)

    borders = OxmlElement('w:tblBorders')

    # 顶线�?.5pt 单线
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '12')  # 12 eighth-points = 1.5pt
    top.set(qn('w:space'), '0')
    top.set(qn('w:color'), '000000')
    borders.append(top)

    # 底线�?.5pt 单线
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), '000000')
    borders.append(bottom)

    # 左线：无
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'none')
    left.set(qn('w:sz'), '0')
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), 'auto')
    borders.append(left)

    # 右线：无
    right = OxmlElement('w:right')
    right.set(qn('w:val'), 'none')
    right.set(qn('w:sz'), '0')
    right.set(qn('w:space'), '0')
    right.set(qn('w:color'), 'auto')
    borders.append(right)

    # 内横线：�?    insideH = OxmlElement('w:insideH')
    insideH.set(qn('w:val'), 'none')
    insideH.set(qn('w:sz'), '0')
    insideH.set(qn('w:space'), '0')
    insideH.set(qn('w:color'), 'auto')
    borders.append(insideH)

    # 内竖线：�?    insideV = OxmlElement('w:insideV')
    insideV.set(qn('w:val'), 'none')
    insideV.set(qn('w:sz'), '0')
    insideV.set(qn('w:space'), '0')
    insideV.set(qn('w:color'), 'auto')
    borders.append(insideV)

    tblPr.append(borders)

    # 表头行下方添�?0.75pt 分隔线（通过为表头�?单元格�?�?��边�?�?    if len(table.rows) > 0:
        for cell in table.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            header_bottom = OxmlElement('w:bottom')
            header_bottom.set(qn('w:val'), 'single')
            header_bottom.set(qn('w:sz'), '6')  # 6 eighth-points = 0.75pt
            header_bottom.set(qn('w:space'), '0')
            header_bottom.set(qn('w:color'), '000000')
            tcBorders.append(header_bottom)
            # 清除单元格其它边�?            for side in ['top', 'left', 'right', 'insideH', 'insideV']:
                s = OxmlElement(f'w:{side}')
                s.set(qn('w:val'), 'none')
                s.set(qn('w:sz'), '0')
                s.set(qn('w:space'), '0')
                s.set(qn('w:color'), 'auto')
                tcBorders.append(s)
            tcPr.append(tcBorders)


def add_three_line_table(headers, rows, caption=None, col_widths=None):
    """添加三线�?���?��序与表�?�?""
    # 表�?（表�?+ 表�?文字�?    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(6)
        cap.paragraph_format.space_after = Pt(4)
        run_num = cap.add_run(f'{_next_table_num()}  ')
        run_num.font.name = '黑体'
        run_num.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run_num.font.size = Pt(10)
        run_num.bold = True
        run_txt = cap.add_run(caption)
        run_txt.font.name = '宋体'
        run_txt.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run_txt.font.size = Pt(10)

    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.autofit = True

    # �?��表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10)
        run.bold = True

    # �?��数据
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(10)

    # 应用三线表格�?    set_three_line_table(table)

    doc.add_paragraph()  # 表后空�?
    return table


def add_image(img_filename, width_inches=5.0, caption=None):
    """插入图片（含图序与图题）"""
    img_path = os.path.join(IMG_DIR, img_filename)
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_inches))
        # 图�?（图�?+ 图�?文字�?        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(6)
            run_num = cap.add_run(f'{_next_figure_num()}  ')
            run_num.font.name = '黑体'
            run_num.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run_num.font.size = Pt(10)
            run_num.bold = True
            run_txt = cap.add_run(caption)
            run_txt.font.name = '宋体'
            run_txt.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run_txt.font.size = Pt(10)
    else:
        add_para(f'[图片缺失：{img_filename}，�?手动插入]', indent=False)


# ══════════════════════════════════════════════════�?# 封面
# ══════════════════════════════════════════════════�?
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('�?�?���?)
title_run.font.name = '黑体'
title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
title_run.font.size = Pt(26)
title_run.bold = True

doc.add_paragraph()

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run('专项赛一  基于ROS的单臂机器人\n—�?�俄罗斯方块�?��化拼接系�?)
sub_run.font.name = '黑体'
sub_run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
sub_run.font.size = Pt(18)

for _ in range(4):
    doc.add_paragraph()

for item in ['项目团队：�?汉大�? 创意�?�?, '提交日期�?026�?�?0�?, '版本号：4.0']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(item)
    r.font.name = '宋体'
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(14)

doc.add_page_break()

# ══════════════════════════════════════════════════�?# �?��
# ══════════════════════════════════════════════════�?
add_heading_styled('�?��', level=1)
add_toc()
doc.add_page_break()

# ══════════════════════════════════════════════════�?# 摘�?
# ══════════════════════════════════════════════════�?
add_heading_styled('摘�?', level=1)

add_para(
    '�?���?��档面向�?九届�?��高校智能机器人创意大赛专项赛�?（基于ROS的单臂机器人—�?�俄罗斯方块�?��化拼接）�?
    '系统阐述了一套基于机器人操作系统（ROS）�?架的智能单臂机器人解决方案�?�系统以睿尔曼RM65-B�?��协作机�?臂为核心执�?单元�?
    '�?ntel RealSense D435深度相机为�?觉感知前�?��通过YOLOv8实例分割模型实现俄罗�?��块的实时�?测与位姿估�?�?
    '结合Pierre Dellacherie�?��式评估与深度Q网络（DQN）混合决策算法生成最优拼接策略，'
    '并利用MoveIt!运动规划框架与自定义轨迹优化器完成高精度拾取-放置操作�?
)

add_para(
    '文档从系统架构�?计�?�硬件�?�型论证、软件平台构建�?��?觉标定与识别算法、拼接决策模型�?�轨迹�?划与控制策略�?��维度展开详细论述�?
    '实验数据表明，系统在标准化测试环境中达到93.5%的方块识�?��度（mAP@0.5），�??定位�?��控制�??.02 mm以内�?
    '单�?拾取-放置�?��时间�?.2秒，连续运�?成功�?5.3%。针对当前存在的视�?处理延迟�?0 ms/帧）与动态环境�?�应性不足等�??�?
    '�?��提出了基于TensorRT模型量化加�?��?�自适应光照补偿与�?传感器融合的优化方�?，为后续系统�?��提供了明�?���?��径�??
)

add_para('关键词：ROS；单臂机器人；俄罗斯方块；YOLOv8；深度强化�?习；运动规划；手眼标�?, bold=True)

doc.add_page_break()

# ══════════════════════════════════════════════════�?# 1. 系统概述
# ══════════════════════════════════════════════════�?
set_chapter(1)
add_heading_styled('1. 系统概述', level=1)

add_heading_styled('1.1 项目背景与竞赛任务分�?, level=2)

add_para(
    '�?��届中国高校智能机器人创意大赛专项赛一要求参赛队伍基于ROS框架，�?计并实现�?套单臂机器人系统�?
    '能�?�?��完成俄罗�?��块的视�?识别、智能拼接与高效堆叠。竞赛任务的核心�?�?��战涵盖三�?���?��'
    '(1) 非结构化�??下的实时视�?感知—�?��?多类�??��?姿�?�的俄罗�?��块进行鲁棒�?测与精确位姿估�?�?
    '(2) 有限工作空间内的�?优决策�?��?��?计高效的拼接算法以最大化空间利用率与堆叠稳定性；'
    '(3) 实时约束下的高精度运动控制�?��?�在保证安全性的前提下实现从决策到执行的�??�闭�???
    '上述三项子任务分�??应�?算机视�?、组合优化与机器人运动�?划三�?��心研究�?域，'
    '要求系统在感�?决策-执�?三个层面实现深度协同�?
)

add_para(
    '�?��案采�?视�?感知—智能决策�?�运动执�?三层递阶架构，各层之间�?�过ROS话�?/服务机制实现松�?�合通信�?
    '视�?感知层基于YOLOv8深度学习模型与手眼标定模块，完成方块类别识别与三维空间定位；'
    '智能决策层融合Pierre Dellacherie�?��式评估与深度Q网络，动态生成最优放�?��略；'
    '运动执�?层依托MoveIt!规划框架与RM65-B机�?臂驱�?��实现从关节空间到笛卡尔空间的高精度轨迹跟�???
)

add_heading_styled('1.2 系统总体架构', level=2)

add_para(
    '系统采用分层架构设�?，自底向上划分为�?��抽象层�?�驱动�?�信层�?�核心算法层与应用层四个层级�?
)

add_para(
    '�?��抽象层：包含RM65-B�?��机�?臂本体�?�Intel RealSense D435深度相机、气动吸盘末�?��行器�?
    'NVIDIA Jetson边缘计算单元、Arduino串口控制模块及工业级供电与�?�信模块，为上层提供统一的�?备能力接口�??
)

add_para(
    '驱动通信层：基于ROS Noetic分布式架构，通过rm_driver节点实现CAN总线�? Mbps）与机�?臂伺服驱动器的实时�?�信�?
    '通过realsense2_camera节点封�?RealSense SDK 2.0的彩�?深度数据流�?�各节点间采用话题（Topic）机制进行异步数�?��发，'
    '采用服务（Service）机制完成同步指令传输，核心控制指令延迟低于10 ms�?
)

add_para(
    '核心算法层：部署视�?识别、拼接决策�?�轨迹�?划与运动控制四大算法模块，各模块封�?为独立ROS节点�?
    '视�?识别模块（yolo_detector节点）输出方块类�??�质心坐标与旋转角度；拼接决策模块（tetris_planner节点�?
    '计算�?优放�?���?��轨迹规划模块（trajectory_planner节点）生成无碰撞关节轨迹�?
    '运动控制模块（motion_controller节点）负责轨迹执行与异常处理�?
)

add_para(
    '应用层：集成rviz三维�??化�?�rqt实时监控面板与自定义调试界面，为操作者提供系统状态全�??图，'
    '�?��任务�?��、参数在线调整与故障诊断�?
)

add_para('系统运�?流程如图1-1�?示�??, indent=True)
add_image('img0.jpeg', 5.5, caption='系统运�?流程�?)

add_heading_styled('1.3 关键�?�?���?, level=2)

add_para('系统在�?计阶段确定的关键�?�?��标与实测值�?�?-1�?示�??, indent=True)

add_three_line_table(
    ['指标类别', '指标�?, '�?���?, '实测�?],
    [
        ['视�?识别', '方块�?测mAP@0.5', '�?90%', '93.5%'],
        ['视�?识别', '单帧推理时间', '< 60 ms', '50 ms'],
        ['视�?识别', '位姿估�?角度�?��', '< 3°', '1.8°'],
        ['运动控制', '重�?定位精度', '�?±0.05 mm', '±0.02 mm'],
        ['运动控制', '�??轨迹跟踪�?��', '< 1 mm', '0.8 mm'],
        ['任务效率', '单�?拾取-放置周期', '< 10 s', '8.2 s'],
        ['任务效率', '连续运�?成功�?, '�?90%', '95.3%'],
        ['系统�?���?, '平均无故障运行时�?, '�?2 h', '3.5 h'],
    ],
    caption='系统设�?关键�?�?���?
)

doc.add_page_break()

# ══════════════════════════════════════════════════�?# 2. �?��与软件平�?# ══════════════════════════════════════════════════�?
set_chapter(2)
add_heading_styled('2. �?��与软件平�?, level=1)

add_heading_styled('2.1 �?��配置与�?�型论证', level=2)

add_para(
    '�?��系统的�?�型围绕竞赛规则约束（臂展≤1000 mm、供电电压≤220 V、电流≤10 A）与任务�?求（有效负载�?00 g�?
    '亚�?米级定位精度、实时�?觉�?理能力）展开，各核心组件的�?�型依据与技�?��数�?述�?下�??
)

add_heading_styled('2.1.1 机�?臂系�?, level=3)

add_para(
    '执�?机构选用睿尔曼RM65-B�?��协作机�?臂�?��?�型论证如下�?
    '(1) 工作半径610 mm，在规则限定�?000 mm臂展范围内最大化工作空间覆盖�?
    '(2) 重�?定位精度±0.02 mm，优于任务所�?的�?.05 mm精度阈�?�，为高精度抓取提供充分裕量�?
    '(3) �?大负�? kg（峰�? kg），远超俄罗�?��块（单块�?0 g）与�??执�?�?��0.3 kg）的负载�?求，'
    '保证加减速阶段不会触发过载保护�??
)

add_para(
    '机�?臂机�?��用高强度�?��铝合金（7075-T6）制造，�?���?.2 kg（含控制�?��，刚�?重量比优异�??
    '�?��关节均配置Maxon EC-4pole无刷直流伺服电机（�?定功�?50 W/关节），'
    '配合14-bit绝�?式�?编码器实�?.022°的�?度分辨率，关节力矩传感器（量�?-50 Nm，响应时�?5 ms�?
    '为力控柔顺操作与碰撞�?测提供硬件基�?。关节构型为肩部2�?��度（�?��±90°、偏�??80°）�??
    '肘部1�?��度（�?��±120°）�?�腕�?�?��度（�?��±90°、偏�??80°、翻滚�?60°），'
    '该六�?��度构型在工作空间内不存在奇异位姿盲区，满足俄罗斯方块多�?度拾取与放置的灵活�?��?求�??
)

add_para(
    '�??执�?器为定制气动吸盘，主体材料为碳纤维�?合材料（密度1.6 g/cm³），�?��0.3 kg�?
    '有效抓取范围0-80 mm。气动回�?���?��空发生器（最大真空度-85 kPa）与电�?换向�?（响应时�?10 ms），'
    '配合0.5 MPa压缩气源�?��生约20 N吸附力，�?0-50 mm尺�?的方块形成可靠抓取�??
    '吸盘接触面�?有�?橡胶缓冲层，在保护方块表面的同时增强摩擦系数（μ≈0.8）�??
)

add_heading_styled('2.1.2 视�?感知系统', level=3)

add_para(
    '视�?输入设�?选用Intel RealSense D435深度相机。�?相机基于主动立体红�?（Active Stereo IR）技�?��'
    '通过左右红�?相机（基�?0 mm）与红�?点阵投影器协同工作，�?.3-3 m范围内生成�?密深度图�?
    '核心参数：RGB传感器分辨率1920×1080@30 fps（�?场�?69.4°×42.5°），深度传感器分辨率1280×720@30 fps（�?场�?87°×58°），'
    '深度精度±2%@1 m。相较于结构光方案（如RealSense D415），D435�?.5-1.5 m典型工作距�?下具有更宽�?场�?�?
    '更�?�合覆盖竞赛�? m²的工作台�??�内置Bosch BMI055�?��IMU（加速度±2g，�?速度±2000°/s）用于动态姿态补偿�??
)

add_para(
    '相机采用"眼在手�?"（Eye-to-Hand）安装方式，固定于工作台�?��方约0.8 m处，�?��角约60°�?
    '使相机�?场完整�?盖工作区域�?��?安�?方式使相机坐标系相�?于基座坐标系保持固定�?
    '�?化了手眼标定过程，同时避免了"眼在手上"（Eye-in-Hand）方案中因机械臂运动导致的图像模糊问题�??
)

add_heading_styled('2.1.3 计算与�?�信系统', level=3)

add_para(
    '主控计算机配置Intel Core i7-9700处理�?��8�?线程，基�?频率3.0 GHz，最大睿�?.7 GHz）�??
    '16 GB DDR4-2666 RAM与NVIDIA GeForce RTX 2080 Ti GPU�?1 GB GDDR6�?352 CUDA核心），'
    '为YOLOv8模型推理与DQN�?��提供充足算力保障。边缘�?部署NVIDIA Jetson Nano'
    '�?核ARM Cortex-A57�? GB LPDDR4�?28核Maxwell GPU，功�?0 W）作为辅助�?算节点，'
    '分担图像预�?理与串口通信任务�?
)

add_para(
    '系统内部通信采用千兆工业以太网交换机（带�? Gbps，转发延�?5 μs），通过TCP/UDP协�?承载ROS节点间数�?���???
    '机�?臂与主控之间通过CAN总线（波特率1 Mbps）传输关节指令与状�?�反馈，物理层采用双绞屏蔽线保证抗干扰能力�??
    'Arduino Uno通过USB串口�?dev/ttyUSB0，波特率9600）连接电磁阀继电器模块，控制吸盘的吸�?释放动作�?
)

add_para(
    '供电系统采用明纬LRS-350-24工业级开关电源（输出24 V/14.6 A，�?定功�?50 W），'
    '输入电压范围100-240 V AC（兼容竞赛场地供电），输出纹�?50 mV，效率≥89%�?
    '为机械臂控制器�?�相机与交换机统�?供电。系统主要硬件组件配�?��总�?�?-1�?示�??
)

add_three_line_table(
    ['类别', '型号/规格', '关键参数', '用�??],
    [
        ['机�?�?, '睿尔�?RM65-B', '6-DoF, 610 mm臂展, ±0.02 mm精度', '核心执�?单元'],
        ['深度相机', 'Intel RealSense D435', '1920×1080 RGB, 1280×720 Depth, 30 fps', '视�?感知'],
        ['�??执�?�?, '定制气动吸盘', '碳纤�? 0.3 kg, 20 N抓取�?, '方块拾取'],
        ['主控计算�?, '�?��装工作站', 'i7-9700, 16 GB RAM, RTX 2080 Ti', '核心计算'],
        ['边缘计算', 'NVIDIA Jetson Nano', '4 GB LPDDR4, 128核GPU, 10 W', '辅助计算'],
        ['交换�?, '工业级千兆以�?��', '1 Gbps, <5 μs延迟', 'ROS节点通信'],
        ['电源', '明纬 LRS-350-24', '24 V/14.6 A, 350 W, 纹波<50 mV', '系统供电'],
        ['串口控制', 'Arduino Uno', 'ATmega328P, USB-UART, 9600 bps', '吸盘通断控制'],
    ],
    caption='系统主�?�?��组件配置'
)

add_heading_styled('2.2 �?��平台', level=2)

add_heading_styled('2.2.1 操作系统与ROS框架', level=3)

add_para(
    '�?��系统运�?于Ubuntu 20.04 LTS（Linux内核5.4.0），选用ROS Noetic Ninjemys发�?版�??
    'ROS提供分布式的节点间�?�信基�?设施，其核心机制包括：基于TCP/UDP的话题（Topic）发�?订阅模型�?
    '适用于传感器数据流等高带宽单向传输场�?��基于请求-应答的服务（Service）模型，适用于机械臂指令下发等同步操作；'
    '参数服务�?��Parameter Server）用于存储标定矩阵�?��?�度限制等配�?���???
)

add_para(
    '安�?的ROS功能包包�?��ros-noetic-desktop-full元包（涵盖rviz、rqt、tf2等核心工具）�?
    'ros-noetic-moveit全�?运动规划框架、ros-noetic-realsense2-camera相机驱动�?
    'ros-noetic-cv-bridge与image-transport图像传输工具链�?�自定义功能包rm_robot'
    '（含rm_driver、rm_msgs、rm_65_moveit_config、rm_65_demo等子包）由睿尔曼官方提供�?
    '封�?了机械臂底层驱动与MoveIt!配置�?
)

add_heading_styled('2.2.2 核心�?��依赖', level=3)

add_para(
    '视�?处理栈：OpenCV 4.5.0（启用SSE4.2与AVX2指令集加速）、Intel RealSense SDK 2.50.0�?
    'YOLOv8（Ultralytics发�?版，基于PyTorch 1.13）�?�CUDA 11.6与cuDNN 8.4�?
)

add_para(
    '运动规划栈：MoveIt! 1.1.7（集成OMPL 1.5.2规划库与TRAC-IK逆运动�?求解�?���?
    'rm_msgs�?��义消�?��（定义Arm_Current_State、MoveJ_P_Cmd等消�?��型，传输频率500 Hz）�??
)

add_para(
    '�?发与运维工具：Visual Studio Code（集成ROS与Python插件）�?�Git 2.25（遵循GitFlow分支管理规范）�??
    'rosdep依赖管理工具�?
)

add_three_line_table(
    ['�?��组件', '版本', '用�??, '备注'],
    [
        ['Ubuntu', '20.04 LTS (Kernel 5.4)', '操作系统', '长期�?���?025�?],
        ['ROS', 'Noetic Ninjemys', '机器人中间件', '集成MoveIt!、rviz�?],
        ['OpenCV', '4.5.0', '图像处理�?, 'SSE4.2/AVX2优化'],
        ['YOLOv8', 'Ultralytics 8.0+', '�?���?测模�?, 'PyTorch 1.13后�?'],
        ['RealSense SDK', '2.50.0', '深度相机驱动', '�?��校准工具'],
        ['MoveIt!', '1.1.7', '运动规划框架', 'OMPL 1.5.2 + TRAC-IK'],
        ['CUDA/cuDNN', '11.6 / 8.4', 'GPU计算平台', 'NVIDIA RTX 2080 Ti'],
        ['PyTorch', '1.13.0', '深度学习框架', 'CUDA编译'],
    ],
    caption="核心�?���??与依赖版�?
)

add_heading_styled('2.3 ROS通信架构', level=2)

add_para(
    '系统运�?时，ROS主节点（roscore）在11311�?��监听，�?理所有节点的注册与�?�信配�?。核心数�?��如下�?
)

add_para(
    '(1) 图像数据流：realsense2_camera节点�?0 Hz频率发布/camera/color/image_raw话�?（sensor_msgs/Image），'
    'yolo_detector节点订阅后进行推理，�?测结果以�?��义消�?��布至/yolo_ros/detections话�?�?
)

add_para(
    '(2) 机�?臂状态流：rm_driver节点�?00 Hz频率发布/rm_driver/Arm_Current_State话�?（rm_msgs/ArmState），'
    '提供关节角度、末�?��姿与力矩传感器�?数�?�transform节点订阅该话题以获取手眼标定�?�?的实时位姿�??
)

add_para(
    '(3) 控制指令流：motion_controller节点发布/rm_driver/MoveJ_P_Cmd话�?（rm_msgs/MoveJ_P），'
    'rm_driver节点接收后�?�过CAN总线�?��至伺服驱动器执�?。指令发布�?率约2 Hz（受rospy.sleep间隔限制）�??
)

add_para(
    '(4) 吸盘控制流：xipan/fangzhi节点通过serial协�?�?dev/ttyUSB0, 9600 bps�?
    '向Arduino发�?�单字节指令�?x01吸取/0x00释放），Arduino驱动电�?�?完成气路切换�?
)

add_para(
    '(5) �??化与调试流：rviz订阅/joint_states�?rm_driver/Arm_Current_State话�?获取关节数据�?
    '渲染机�?臂三维模型与点云叠加；rqt_plot绘制关节角度/力矩随时间变化曲线；'
    'rqt_console汇�?�各节点日志（按INFO/WARN/ERROR分级过滤）�??
)

doc.add_page_break()

# ══════════════════════════════════════════════════�?# 3. �?发与操作流程
# ══════════════════════════════════════════════════�?
set_chapter(3)
add_heading_styled('3. �?发与操作流程', level=1)

add_heading_styled('3.1 �??�?��与编�?, level=2)

add_para(
    '系统工作空间（workspace）位于~/ws_rmrobot/�?��，采用标准catkin工作空间结构�?
)

add_code_block(
    '~/ws_rmrobot/\n'
    '├─�? src/                    # 源代码目录\n'
    '�?  ├─�? rm_robot/           # 睿尔曼官方功能包\n'
    '�?  �?  ├─�? rm_driver/      # 机�?臂驱动节点\n'
    '�?  �?  ├─�? rm_msgs/        # �?��义消�?��服务定义\n'
    '�?  �?  ├─�? rm_description/ # URDF模型文件\n'
    '�?  �?  ├─�? rm_65_moveit_config/  # MoveIt!配置包\n'
    '�?  �?  └─�? rm_65_demo/     # 应用示例脚本\n'
    '�?  ├─�? yolo_ros/           # YOLOv8 ROS封�?\n'
    '�?  └─�? tetris_planner/     # 拼接决策算法包\n'
    '├─�? build/                  # 编译�?��文件\n'
    '├─�? devel/                  # �?发环境变量脚本\n'
    '└─�? install/                # �?��署安装目录\n'
)

add_para(
    '编译流程�?1) 在CMakeLists.txt�?��加依赖项find_package(catkin REQUIRED COMPONENTS roscpp rospy '
    'moveit_ros_planning realsense2_camera)，执行catkin build --cmake-args -DCMAKE_BUILD_TYPE=Release'
    '进�?Release模式编译（启�?O3优化与NDEBUG宏），编译�?�时�?0-15分钟（取决于CPU核心数）�?
    '(2) 编译完成后执行source devel/setup.bash初�?化ROS�??变量�?
    '(3) 运�?rosversion -d验证ROS发�?版标识为noetic。若遇依赖缺失，运�?'
    'rosdep install --from-paths src --ignore-src -r -y�?��解析并安装缺失的系统依赖（�?librealsense2-dev）�??
)

add_heading_styled('3.2 系统初�?�?, level=2)

add_para('系统�?��序列严格按照以下顺序执�?，以�?��依赖关系正确�?)

add_para(
    'Step 1 �?�?��ROS主节点：roscore &（后台运行），roscore�?��耗时<5秒，监听�?��11311�?
    '�?��日志�?��记录至~/.ros/log/�?��，文件名以启动时间戳命名�?
)

add_para(
    'Step 2 �?�?��机�?臂控制节点：roslaunch rm_control rm_control.launch，�?launch文件加载rm_driver节点�?
    '加载URDF模型（robot_description参数），初�?化CAN总线通信并完成机械臂零点标定，�?�时�?0秒�??
    '�??�过rosnode list验证rm_driver_node�?��成功注册�?
)

add_para(
    'Step 3 �?�?��相机节点：roslaunch realsense2_camera rs_camera.launch，加载内参标定文件（calibration.yaml），'
    '�?��彩色与深度数�?��（均�?0 Hz发布），同时�?活点云生成（/camera/depth/color/points话�?）�??
)

add_para(
    'Step 4 �?�?��MoveIt!运动规划框架：roslaunch rm_65_moveit_config demo.launch，加载运动�?求解�?��TRAC-IK）�??
    '碰撞�?测场�?��Octomap）与rviz�??化界�???
)

add_heading_styled('3.3 驱动与运动控�?, level=2)

add_para(
    '机�?臂驱动层通过rm_driver节点实现与伺服控制器的CAN总线通信。CAN帧采用标准格式（11-bit标识符）�?
    '数据�?字节，传输波特率1 Mbps，指令周�? ms。驱动层向上层暴露两�?��心接口：'
)

add_para(
    '(1) 话�?接口：�?�?rm_driver/MoveJ_P_Cmd话�?（消�?��型rm_msgs/MoveJ_P），'
    '每条消息包含�?���??位姿（x, y, z, roll, pitch, yaw）与运动速度比例（speed∈[0, 1]），'
    '驱动层内部完成�?�运动�?解算与关节空间插值，生成平滑轨迹后下发至伺服驱动器�??
)

add_para(
    '(2) 状�?�反馈：�?00 Hz频率发布/rm_driver/Arm_Current_State话�?�?
    '消息包含6�?��节的角度（°）、�?速度（�?s）�?�力矩（Nm）与�??执�?器在基座坐标系下�?D位姿�?
)

add_para(
    '应用层运动控制�?�过moveJ_P.py脚本实现。�?脚本通过订阅/rm_driver/Plan_State话�?'
    '（消�?��型rm_msgs/Plan_State，布尔�?�表示轨迹执行是否成功）实现�?��运动同�?�?
    '发布MoveJ_P�?��位姿后，脚本调用wait_for_move()函数�?0 Hz频率�??Plan_State状�?�，'
    '超时阈�?��?�?0秒，任一轨迹点超时或返回失败均�?录WARN级别日志。相较于固定延时等待�?
    '该基于状态反馈的同�?机制能自适应运动耗时波动，同时保证轨迹点间的严格时序约束�?
    '�??姿�?�的roll与pitch分量由机械臂安�?几何关系固定为�? rad�?.005 rad，yaw分量由上层决策模块动态�?算�??
)

add_heading_styled('3.4 实时监控与启�?, level=2)

add_para(
    '系统提供多层次的实时监控手�?�?1) 命令行监控：rostopic echo /rm_driver/Arm_Current_State'
    '实时打印关节状�?�流，rostopic hz /rm_driver/MoveJ_P_Cmd统�?控制指令发布频率�?
    'rosnode ping rm_driver_node�?测节点存活状态；'
    '(2) 图形化监控：rviz显示机�?臂三维模型�?�点云叠加与碰撞�?测场�?��'
    'rqt_plot绘制关节角度/力矩随时间变化曲线，rqt_console汇�?�各节点日志（按INFO/WARN/ERROR分级过滤）；'
    '(3) �?�?���?��通过roslaunch rm_bringup rm_robot.launch同时�?��上述�?有节点（�?0秒完成）�?
    'launch文件通过required="true"属�?�保证关�?��点异常�??出时�?��终�?整个系统�?
)

add_heading_styled('3.5 故障排除指南', level=2)

add_three_line_table(
    ['故障现象', '�?��原因', '排查方法', '解决方�?'],
    [
        ['roscore�?��失败\n（�?口占�?��',
         '11311�?���?��进程占用',
         'lsof -i:11311或ss -tlnp | grep 11311',
         'kill旧进程后重启；或�?��ROS_MASTER_URI'],
        ['机�?臂无响应\n（CAN通信�?���?,
         'CAN线缆松动或USB转CAN\n适配器掉�?,
         'dmesg | grep can�?查内核日志；\ncandump can0监听总线数据',
         '重新插拔USB线缆；\n重启rm_driver节点'],
        ['相机无数�?��',
         'USB 3.0带�?不足或\n线缆质量不达�?,
         'rs-enumerate-devices列出设�?；\nlsusb -t�?查连接�?�率（应�? Gbps�?,
         '更换高质量有源USB 3.0延长线；\n降低分辨率或帧率临时缓解'],
        ['MoveIt!规划失败',
         '�?��位姿超出工作空间\n或与场景碰撞',
         '在rviz�?��动拖拽末�?��证\n�?��性；�?�?RDF关节限位',
         '调整�?��位姿；确�?nrobot_description参数正确加载'],
        ['吸盘无法抓取',
         '气压不足或电磁阀\n线路故障',
         '气压表�?数应�?.4 MPa；\n万用表测量电磁阀线圈电阻',
         '�?查气源与管路密封；\n更换电�?�?驱动模块'],
        ['YOLO推理超时',
         'GPU显存不足或\n�?��他进程占�?,
         'nvidia-smi查看GPU利用率\n和显存占�?,
         '终�?无关GPU进程；\n降低输入图像分辨率至1280×720'],
    ],
    caption="常�?故障排查与解决方�?
)

doc.add_page_break()

# ══════════════════════════════════════════════════�?# 4. �?�?��现细�?# ══════════════════════════════════════════════════�?
set_chapter(4)
add_heading_styled('4. �?�?��现细�?, level=1)

# �?�? 4.1 �?�?
add_heading_styled('4.1 视�?标定方法', level=2)

add_para(
    '视�?标定�?��接�?觉感知与机�?臂执行的桥�?�?��，其�?��在于建立像素坐标系�?�相机坐标系与机械臂基座坐标系三者之间的'
    '精确映射关系。标定流程分为相机内参标定�?�深度校正与手眼标定三个阶�?，标定精度直接决定抓取操作的准确性�??
)

add_heading_styled('4.1.1 相机内参标定', level=3)

add_para('相机内参标定基于针孔相机模型，其数�?形式为：')

add_formula(r's \cdot [u, v, 1]^T = K \cdot [R \mid t] \cdot [X, Y, Z, 1]^T')

add_para(
    '其中�?u, v)为像素坐标，(X, Y, Z)为世界坐标，s为尺度因子，K�?×3内参矩阵，[R|t]为�?参矩阵�??
    '内参矩阵K包含焦距(f_x, f_y)与主�?c_x, c_y)四个参数�?
)

add_formula(r'K = \begin{bmatrix} f_x & 0 & c_x \\ 0 & f_y & c_y \\ 0 & 0 & 1 \end{bmatrix}')

add_para(
    '标定实施过程�?1) 使用9×6棋盘格标定板（单�?0 mm，高反差黑白印刷），'
    '从不同�?度（倾斜±30°）与距�?�?.5-1.5 m）采�?0帧图像，�?��棋盘格占�?���?0%以上视场�?
    '(2) 对每帧图像调用cv2.findChessboardCorners()�?测内角点�?
    '再�?�过cv2.cornerSubPix()进�?亚像素精化（终�?精度0.1像素）；'
    '(3) 将所有图像�?点坐标输�?v2.calibrateCamera()�?
    '通过�?小化重投影�?�??��?�目标函数为Σ||p_ij - 投影(P_i, K, dist, R_j, t_j)||²—�??
    '求解内参矩阵K与畸变系数dist=(k�? k�? p�? p�? k�?�?
    '标定结果的重投影�?���?0.3像素，否则需剔除低质量图像并重新标定�?
)

add_heading_styled('4.1.2 深度校�?与点云生�?, level=3)

add_para(
    'RealSense D435深度传感器存在系统�?�的深度测量偏差（主要由温度漂移与红外投影器老化引起），'
    '�?通过Intel RealSense SDK的自校准工具（rs-depth-quality）进行在线校正�??
    '校�?流程：将相机对准平面�?��（�?白�?）距�? m处，运�?�?��准程序，'
    'SDK内部优化深度比例因子与偏移量，使RMS深度�?��降至<2 mm�?
)

add_para('校�?后的深度图�?�过以下�?��逐像素投影为三维点云�?)

add_formula(r'X = (u - c_x) \cdot Z / f_x')
add_formula(r'Y = (v - c_y) \cdot Z / f_y')
add_formula(r'Z = \operatorname{depth}(u, v) / \operatorname{scale}')

add_para(
    '其中depth(u,v)为深度传感器在像�?u,v)处的原�?读数�?6-bit整数），scale为深度比例因子（D435默�?�?.001 m/单位）�??
    '生成的点云�?�过tf2库变换至机�?臂基座坐标系，供后续抓取规划使用�?
)

add_heading_styled('4.1.3 手眼标定', level=3)

add_para(
    '手眼标定求解相机坐标系到机�?臂基座坐标系的刚体变换矩阵ᵇT_c∈SE(3)�?
    '�?���?���?眼在手�?"配置的Tsai-Lenz标定算法，�?算法将标定问题分解为旋转部分与平移部分的顺序求解�?
    '在数值稳定�?�和计算效率之间取得了良好平衡�??
)

add_para(
    '算法核心原理：令A_i表示机�?臂末�?��第i�?��姿到第i+1�?��姿的齐�?变换�?
    'B_i表示标定板在第i帧到第i+1帧的相机观测变换，则手眼变换X满足AX=XB方程�?
)

add_formula(r'A_i \cdot X = X \cdot B_i \quad (i = 1, 2, \dots, n-1)')

add_para(
    '将A_i和X分解为旋�?��分和平移部分：A_i=[R_Ai, t_Ai; 0, 1], X=[R_X, t_X; 0, 1], B_i=[R_Bi, t_Bi; 0, 1]�?
    '则原方程展开为：'
)

add_formula(r'R_{A_i} \cdot R_X = R_X \cdot R_{B_i}')
add_formula(r'R_{A_i} \cdot t_X + t_{A_i} = R_X \cdot t_{B_i} + t_X')

add_para(
    'Tsai-Lenz算法求解过程分为两�?�?1) 旋转部分—�?�利用Rodrigues�?��将R_Ai和R_Bi分别�?��为旋�?��量，'
    '通过�?小二乘法求解R_X的旋�?��量表示，再�?�过SVD分解强制正交化以保证R_X∈SO(3)�?
    '(2) 平移部分—�?�将已求得的R_X代入平移方程，构建超定线性方程组(I-R_Ai)·t_X=t_Ai-R_X·t_Bi�?
    '通过�??�法求解�?小二乘解。标定过程采集不少于15组机械臂位姿-标定板图像�?�?
    '位姿覆盖工作空间80%以上范围，最终标定结果的平移�?��<1 mm，旋�??�?0.5°�?
)

add_para(
    '标定结果�?×4齐�?变换矩阵形式存储于transformation.yaml配置文件�?
    'transform.py脚本在运行时加载该矩阵，将相机坐标系下的�?��坐标P_cam经由ᵇT_c·P_cam变换至基座坐标系P_base�?
    '系统设�?了动态标定验证机制：每完�?0次拾�?放置�?��后，�?��执�?�?次�?盘格验证�?
    '若平移偏�?1 mm则触发重标定告�?�?
)

add_heading_styled('4.1.4 标定�?��分析与补�?, level=3)

add_para(
    '标定过程�?��主�?�?��源包�?��(1) 角点�?测噪声（亚像素级，σ≈0.1 px）；'
    '(2) 机�?臂绝对定位�?�?��RM65-B的绝对精度约±0.5 mm，主要由关节间隙与连杆弹性变形引起）�?
    '(3) 深度传感器系统偏�?��约�?% @ 1 m，随工作距�?线�?��?大）�?
)

add_para(
    '针�?上述�?��源，系统采用以下补偿策略�?1) 鲁�?回归—�?��?15组标定数�?��加M-估�?�?��Huber损失函数），'
    '降低离群位姿的权重；(2) Kalman滤波在线校�?—�?�以高�?�?00 Hz）IMU数据为过程模型输入，'
    '以低频（10 Hz）�?觉重投影�?��为�?测，实时估�?并补偿相机姿态漂移；'
    '(3) 温度补偿—�?�在机�?臂连�?��行超�?小时后，�?��触发零漂补偿（将各关节回零并重新标定当前位置）�??
)

doc.add_page_break()

# �?�? 4.2 �?�?
add_heading_styled('4.2 俄罗�?��块�?觉识�?, level=2)

add_heading_styled('4.2.1 �?测模型架�?, level=3)

add_para(
    '视�?识别模块基于YOLOv8实例分割模型实现。YOLOv8采用经典的Backbone-Neck-Head三阶段架构：'
)

add_para(
    'Backbone（特征提取网络）：以CSPDarknet-53为�?干，通过跨阶段部分连接（Cross Stage Partial, CSP�?
    '减少�?��重�?计算，在保持特征表达能力的条件下降低�?0%的�?算量。输�?GB图像尺�?1920×1080�?
    '�?次下采样（stride=2）生成�?尺度特征图{C3, C4, C5}，尺度分�?��输入�?/8�?/16�?/32�?
)

add_para(
    'Neck（特征融合网络）：采用改进的PANet（Path Aggregation Network）结构，在FPN�?��向下�?��之�?增加'
    '�?��向上的特征金字�?�?��，使高层�?��信息与底层空间细节实现双向融合�??
    'CSPBlock�?��标准卷积�?���?��C2f模块（CSP Bottleneck with 2 Convolutions），'
    '通过split-merge机制进一步减少参数量�?
)

add_para(
    'Head（�?测头）：采用Anchor-Free的解耦�?测头，将分类任务与回归任务分配给不同的卷�?���???
    '分类分支输出H×W×C的特征图（C=7，�?应O/L/T/Z/z/I七类方块和背�?���?
    '回归分支输出H×W×4的边界�?偏移量�?�相较于Anchor-Based方法，Anchor-Free设�?'
    '消除了anchor超参数的手工设置，在俄罗�?��块尺寸差异显著（20-50 mm）的场景�?��现出更好的泛化�?��??
)

add_para(
    '损失函数采用CIoU Loss（Complete IoU Loss）：'
)

add_formula(r'\mathcal{L}_{\text{CIoU}} = 1 - \text{IoU} + \frac{\rho^2(b, b_{gt})}{c^2} + \alpha \cdot v')

add_para(
    '其中ρ(·)表示预测框与真�?��?�?��点的欧氏距�?，c为最小�?接矩形的对�?线长度，'
    'α为权衡参数，v用于惩罚宽高比不�?致�?�相较于传统IoU Loss，CIoU额�?考虑了中心距离与宽高比，'
    '收敛速度更快。分类损失采用二元交叉熵（BCE Loss），�?��总损失为L=λ_box·L_CIoU+λ_cls·L_BCE'
    '（λ_box=7.5, λ_cls=0.5）�??
)

add_heading_styled('4.2.2 数据集构建与�?��', level=3)

add_para(
    '�?��数据集构建：(1) 在实际竞赛台�?��采集2,500张标注图像，涵盖7类方块�?种旋�??种光照条�?
    '�?00 lux/500 lux/1000 lux）；(2) 通过数据增强生成1,500张合成图像，增强策略包括随机旋转（�?5°）�??
    '�?��调节（�?0%）�?�高�?��声（σ=10）与Cutout�?���?×5网格随机�?��）；'
    '(3) �?��数据集T-LESS�?��部分小物体图像（500张），用于提升模型�?小尺寸目标的�?测鲁棒�?��??
)

add_para(
    '�?��配置：优化器AdamW（lr=0.001, β�?0.9, β�?0.999, weight_decay=0.0005），'
    'Batch Size=16（�?度累�?��，�?�?00 epochs，前3个epoch使用线�??armup�?
    '学习率在150�?40 epoch时衰减为原来�?.1倍�?�数�??处理包括Mosaic拼接（�?�?.5）�??
    '随机HSV变换与尺度缩放（±50%）�?��?练完成后的best.pt模型（FP32精度�?
    '在验证集上达到mAP@0.5=93.5%, mAP@0.5:0.95=78.2%�?
)

add_heading_styled('4.2.3 �?��匹配与姿态精�?, level=3)

add_para(
    'YOLOv8�?测输出的边界框仅提供粗略的位�?���?��为满足抓取所�?的精�?��姿（位置±2 mm，�?度�?°），'
    '系统引入了基于轮廓匹配的姿�?�精化模块（visualize_yolo_degv2.py）�?�其核心流程为：'
)

add_para(
    '(1) �?��提取：从YOLOv8分割掩膜�?��取方块的外轮廓�?边形，使用cv2.findContours()获取有序顶点序列�?
)

add_para(
    '(2) 模板匹配：加载�?定义的七类方块标准轮廓模板（template_contours.npz），每个模板归一化至100�?���?采样点�??
    '对每�??测到的目标轮廓，首先通过等比缩放使模板与�?��的包围盒面积对齐�?
    '随后�?°步长在[0°, 360°)范围内粗搜索�?优旋�??度，再以1°步长在最优粗角度±10°范围内精细搜�?��'
    '搜索的评价指标为�?��IoU（Intersection over Union）：'
)

add_formula(r'\text{IoU}(C_1, C_2) = \frac{\text{Area}(C_1 \cap C_2)}{\text{Area}(C_1 \cup C_2)}')

add_para(
    '(3) 质心计算与偏移补偿：通过图像矩�?算匹配轮廓的质心坐标�?
)

add_formula(r'c_x = \frac{M_{10}}{M_{00}}, \quad c_y = \frac{M_{01}}{M_{00}}')

add_para(
    '其中M_ij=Σ_xΣ_y x^i·y^j·I(x,y)为图像的(i,j)阶空间矩。质心坐标结合深度图对应点的深度值，'
    '经手眼标定矩阵变换至机�?臂基座坐标系，输出为抓取�?��位姿。轮廓匹配IoU阈�?��?�?.85�?
    '低于此阈值的�?测结果�?为�?�?并丢弃�??
)

add_para('视�?识别效果如图4-1和图4-2�?示�??, indent=True)
add_image('img1.png', 4.5, caption='俄罗�?��块�?觉识�?��果（�?��界�?�?)
add_image('img2.jpeg', 4.5, caption='俄罗�?��块�?觉识�?��果（�?��割掩膜）')

doc.add_page_break()

# �?�? 4.3 �?�?
add_heading_styled('4.3 俄罗�?��块拼接算�?, level=2)

add_para(
    '拼接决策�?��统的核心智能模块，其任务�?��给定当前工作台面状�?�与待放�?��块类型的前提下，'
    '输出�?优的放置位置(x, y)与旋�??度θ，以最大化长期�?��收益（清除�?数）�?
    '�?��案采用Pierre Dellacherie�?��式评估与深度Q网络（DQN）相结合的混合决策架构，'
    '兼顾�?��式方法的实时性与深度学习的长期�?划�?�应性�??
)

add_heading_styled('4.3.1 状�?�空间建�?, level=3)

add_para(
    '工作台面离散化为10×20的网格（单格对应实际尺�?20 mm×20 mm），网格状�?�以二�?�矩阵S∈{0,1}^(10×20)表示�?
    '其中S[i][j]=1表示该网格�?方块占据，S[i][j]=0表示空闲�?
    '此�?，当前方块信�?���?��块类别c∈{O, L, l, T, z, Z, I}与初始姿态（质心坐标与旋�??度）�?
    '共同构成DQN的完整状态向量s_t�?
)

add_para(
    '动作空间A定义为所有合法放�?���?��旋转角度的笛卡尔�?��'
)

add_formula(r'\mathcal{A} = \{(x, y, \theta) \mid x \in [0,9], y \in [0,19], \theta \in \{0^\circ, 90^\circ, 180^\circ, 270^\circ\}, \operatorname{is\_valid}(x,y,\theta,c)=\operatorname{True}\}')

add_para(
    '对于任意状�??动作�?s_t, a_t)，状态转移函数T(s_t, a_t)首先将方块c按θ旋�?��放置于网格位�?x, y)�?
    '随后�?查并清除�?有满行（10列均�?的�?），上方�?有�?下移�?��空隙。清除一行获得即时�?励r=+1�?
)

add_heading_styled('4.3.2 Pierre Dellacherie�?��式评�?, level=3)

add_para(
    'Pierre Dellacherie算法通过�?���?��式指标的加权组合为每�??��?�放�?��作评分，'
    '评分函数的�?计体现了对人类专家放�?��略的模仿�?
)

add_formula(r'\operatorname{Score}(s, a) = w_1 \cdot h_{\text{landing}} + w_2 \cdot h_{\text{eroded}} + w_3 \cdot h_{\text{rowTrans}} + w_4 \cdot h_{\text{colTrans}} + w_5 \cdot h_{\text{holes}} + w_6 \cdot h_{\text{wells}}')

add_para(
    '各分量含义与权重如下�?1) Landing Height（着地高度，w�?-4.50）：方块放置后的质心高度，鼓励低处放�?��'
    '(2) Eroded Cells（消除贡�?��w�?+3.42）：该方块贡�?��消除单元格数乘以行转换因子的乘积，鼓励消行；'
    '(3) Row Transitions（�?�?��，w�?-3.22）：行内相邻列的�?占状态变化�?�数，鼓励�?内连�?��'
    '(4) Column Transitions（列�?��，w�?-9.35）：列内相邻行的�?占状态变化�?�数，鼓励列内连�?��'
    '(5) Holes（孔洞数，w�?-7.90）：�?���?��子挡住的空位数量，严格惩罚空洞�?��?�这�?��重绝对�?�最大的指标�?
    '体现�?避免孔洞"在俄罗斯方块策略�?��核心地位�?6) Well Sums（深井惩罚，w�?-3.39）：深度大于1的井槽的�??深度加权和�??
)

add_para(
    '对每�?��放置方块，启发式评估器枚举所有合法放�?��作（通常�?00-500�??��?�）�?
    '选择Score�?高的动作作为�?��式建议a_heuristic。�?算法单�?决策耗时�? ms，满足实时�?��?求，'
    '但缺乏�?�?��方块的长期�?划能力�??
)

add_para('将上述六�?��标代入评分函数，完整展开形式为：', indent=True)
add_formula(r'\operatorname{Score} = -4.50 \cdot H_{\text{land}} - 9.35 \cdot T_{\text{col}} - 7.90 \cdot N_{\text{holes}} - 3.39 \cdot W_{\text{sum}} + 3.42 \cdot E_{\text{cell}} - 3.22 \cdot T_{\text{row}}')
add_para(
    '其中各项依�?为：�?地高�?_land、列�?��数T_col、孔洞�?�数N_holes、深井加权和W_sum�?
    '消除贡献单元格数E_cell、�?�?��数T_row。负权重项为代价（鼓励最小化），正权重项为收益（鼓励�?大化）�??,
    indent=True
)
add_para(
    '算法在枚举所有合法放�?��通常200-500�??��?�）后，通过网格坐标到机械臂基座坐标的线性变换生成实际抓取位姿：',
    indent=True
)
add_formula(r'x_{\text{real}} = x_{\text{grid}} \cdot \Delta s - W/2')
add_formula(r'y_{\text{real}} = y_{\text{grid}} \cdot \Delta s - H/2')
add_para(
    '式中，Δs = 20 mm为单格边长，W = 400 mm、H = 400 mm分别为工作区域�?度与深度�?
    '该变换将离散网格坐标（x_grid∈[0,9], y_grid∈[0,19]）映射至机�?臂基座坐标系（x_real∈[-0.2, 0.2] m, y_real∈[-0.2, 0.2] m）�??,
    indent=True
)

add_heading_styled('4.3.3 深度Q网络（DQN）�?强决�?, level=3)

add_para(
    '为弥补启发式方法的短视�?�缺陷，系统引入DQN模型进�?增强决策。DQN的核心�?�想�??�过深度神经网络'
    '逼近�?优动作价值函数Q*(s, a)，即从状态s出发执�?动作a后，遵循�?优策略所能获得的期望�?��折扣奖励�?
)

add_formula(r"Q^*(s, a) = \mathbb{E}\left[R_t + \gamma \cdot \max_{a'} Q^*(s_{t+1}, a') \mid s_t = s, a_t = a\right]")

add_para(
    '其中γ=0.99为折扣因子，体现了�?�?��奖励的重视程度�?�网络结构�?计：输入层将10×20网格状�?�展平为200维向量，'
    '拼接方块类别one-hot编码�?维）与旋�??度编码（4维），得�?11维输入向量；'
    '隐藏层为3层全连接网络�?12�?56�?28），每层后接BatchNorm与ReLU�?活函数；'
    '输出层为|A|维的Q值向量（经环境过滤后仅保留合法动作的Q值）�?
)

add_para(
    '�?��采用Double DQN与优先经验回放（Prioritized Experience Replay, PER）两项改进：'
)

add_para(
    'Double DQN：将动作选择与动作评估解耦，使用在线网络θ选择动作，使用目标网络θ⁻评估该动作的Q值：'
)

add_formula(r"y_t = r_t + \gamma \cdot Q(s_{t+1}, \arg\max_a Q(s_{t+1}, a; \theta); \theta^{-})")

add_para(
    '这有效缓解了标准DQN中Q值过高估计（overestimation bias）的�??。目标网络θ⁻每C=1000步从在线网络θ复制参数�?
)

add_para(
    '优先经验回放：以TD�?��δ=|y_t-Q(s_t, a_t; θ)|作为样本优先级度量，采样概率P(i)�?δ_i)^α（�?0.6），'
    '通过重�?性采样权重w_i=(N·P(i))^(-β)（β从0.4线�?��?长至1.0）修正优先采样引入的分布偏差�?
    '�?��超参数：学习�?×10⁻⁴，Batch Size=64，经验池容量10⁵条，�?greedy探索策略（ε从1.0线�?��??�?��0.01�?
    '�?�??�?0⁴），在Tetris模拟�??�??�?000个episode�?
)

add_para('�?��过程�?��用的奖励函数定义为：', indent=True)
add_formula(r"R(s, a) = \lambda_{\text{clr}} \cdot N_{\text{clear}}(s') - \lambda_{\text{hgt}} \cdot H_{\text{max}}(s') - \lambda_{\text{hol}} \cdot N_{\text{holes}}(s') - \lambda_{\text{step}}")
add_para(
    '其中，s\'为执行动作a后的后继状�?�，N_clear为清除�?数（每�?+1），H_max为堆叠最大高度，N_holes为新产生的孔洞数�?
    'λ_step = 0.01为每步固定代价以鼓励高效决策。权重�?�?��λ_clr = 10.0（高权重�?励消行）�?
    'λ_hgt = 1.5（抑制高度�?长），λ_hol = 3.0（严格惩罚孔洞）�?
    '该�?计体现了"消除优先、抑制堆�??�避免孔�?的优化目标�??,
    indent=True
)

add_heading_styled('4.3.4 混合决策策略', level=3)

add_para(
    '在线运�?阶�?，系统采用两步混合决策：首先由启发式评估器筛选出Score排名前K=20的�?��?�动作（缩小动作空间），'
    '然后由DQN在这K�??��?�动作中选择Q值最高�?�执行�?��?混合策略兼具以下优势�?
    '(1) �?��式筛选大幅缩减DQN的决策空间，消除大量无效动作（�?�?��放置）的评估�?�?�?
    '(2) DQN在启发式提供的优质�?��?�集�?��行精细�?�择，有效弥补启发式方法缺乏长期规划性的缺陷�?
    '消融实验表明，混合策略相较纯�?��式的平均清除行数提升�?8%，相较纯DQN的决策一致�?�提升约35%�?
)

doc.add_page_break()

# �?�? 4.4 �?�?
add_heading_styled('4.4 机器人轨迹�?�?, level=2)

add_heading_styled('4.4.1 运动学建�?, level=3)

add_para(
    '轨迹规划的基�?�??RM65-B机�?臂进行精�?��正运动�?与�?�运动�?建模�?
    '采用Denavit-Hartenberg（DH）参数法建立�?��由度运动学模型，为每�?��节定义四元组(a_i, α_i, d_i, θ_i)�?
    '相邻连杆的齐次变换矩阵为�?
)

add_formula(r'{}^{i-1}T_i = \operatorname{Rot}(z, \theta_i) \cdot \operatorname{Trans}(0, 0, d_i) \cdot \operatorname{Trans}(a_i, 0, 0) \cdot \operatorname{Rot}(x, \alpha_i)')

add_para(
    '�?�?��杆变换矩阵连乘，得到�??执�?器在基座坐标系下的位姿：'
)

add_formula(r'{}^{0}T_6 = {}^{0}T_1 \cdot {}^{1}T_2 \cdot {}^{2}T_3 \cdot {}^{3}T_4 \cdot {}^{4}T_5 \cdot {}^{5}T_6')

add_para(
    '逆运动�?（IK）求解采用TRAC-IK求解�?��它同时运行基于Newton-Raphson的数值迭代求解器（KDL�?
    '与基于SQP的非线�?�优化求解器，取两�?�中�?优结果�?��?于无解的位姿�?��，TRAC-IK返回欧氏距�?�?近的�??解�??
    'RM65-B�?��关节�?��肩部偏航与腕部偏�?��成冗余自由度，IK求解器利用�?冗余度优化关节运动的平滑性指标：'
)

add_formula(r'J(\Delta\theta) = \frac{1}{2} \Delta\theta^T \cdot W \cdot \Delta\theta \quad \text{s.t.} \quad {}^{0}T_6(\theta+\Delta\theta) = T_{\text{target}}, \quad \theta_{\min} \leq \theta+\Delta\theta \leq \theta_{\max}')

add_para(
    '其中W为�?角权重矩阵，通过赋予大关节较大权重（肩部>肘部>腕部）实�?大关节少动�?�小关节精调"的节能运动策略�??
)

add_heading_styled('4.4.2 分�?轨迹生成', level=3)

add_para(
    '轨迹规划模块（trajectory_planner节点）根�?��取与放置任务的特点，'
    '将每次操作拆分为7�?��迹�?（�?�?-1），每�?采用�?��速度曲线（加�?�?�?减�?�）�?
)

add_three_line_table(
    ['轨迹�?, '运动描述', '起�?位姿', '�?��位姿', '速度�?, '标称时间'],
    [
        ['T1 趋近', '移动至方块\n上方安全高度', '当前位置', '(x_b, y_b, 0.15, 0°)', '0.3', '~1.5 s'],
        ['T2 下降', '垂直下降至\n拾取高度', 'T1终点', '(x_b, y_b, 0.12, 0°)', '0.1', '~0.8 s'],
        ['T3 抓取', '执�?吸盘吸取', '�?, '�?, '�?, '1.6 s'],
        ['T4 �?��', '垂直�?��至\n安全高度', 'T2终点', '(x_b, y_b, 0.15, 0°)', '0.3', '~0.8 s'],
        ['T5 运输', '水平移动至\n放置位上�?, 'T4终点', '(x_p, y_p, 0.15, θ_p)', '0.3', '~1.5 s'],
        ['T6 放置', '垂直下降至\n放置高度', 'T5终点', '(x_p, y_p, 0.13, θ_p)', '0.1', '~0.8 s'],
        ['T7 释放', '执�?释放\n并抬�?, 'T6终点', '(x_p, y_p, 0.15, θ_p)', '0.3', '~1.2 s'],
    ],
    caption="拾取-放置任务分�?轨迹规划参数"
)

add_para(
    '单�?完整拾取-放置�?��的标称周期约8.2秒�?�为减少机�?臂软管缠绕，系统对方块拾取顺序进行了优化�?
    '各循�?��先�?�择x坐标�?小的方块作为拾取�?��，使得末�?��行器的水平位移量在各�?��间趋于均衡，'
    '避免大幅�?返移动�?致的管线�?���?
)

add_para('轨迹规划�?��及的网格坐标到机械臂基座坐标的线性变换关系为�?, indent=True)
add_formula(r'x_{\text{base}} = x_{\text{grid}} \cdot \Delta s - W/2')
add_formula(r'y_{\text{base}} = y_{\text{grid}} \cdot \Delta s - H/2')
add_para(
    '其中Δs = 20 mm为网格单元边长（对应单个方块尺�?），W = H = 400 mm为工作台面的宽度与深度�??
    '该变换将拼接算法输出的�?散网格坐标映射至机�?臂可执�?的三维笛卡尔坐标�?
    'z轴坐标由轨迹规划器根�?��务阶段（趋近/拾取/运输/放置）动态指定（0.12-0.15 m）�??,
    indent=True
)

add_heading_styled('4.4.3 碰撞�?测与避障', level=3)

add_para(
    '轨迹安全性由MoveIt!集成的FCL（Flexible Collision Library）�?撞�?测引擎保障�??
    '工作场景的�?撞模型包�?��(1) 机�?臂自碰撞模型（URDF�?��义的关节与连杆包围盒）；'
    '(2) 工作台面平面约束（z�?的平面�?撞体）；(3) 已放�?��块堆的Octomap占用栅格（分辨率5 mm），'
    '由深度相机的点云数据实时更新。轨迹�?划器在OMPL的RRT-Connect算法�?��成�?撞�?测回调，'
    '每�?采样新节点时即时验证有效性，�?��生成的整条轨迹无碰撞�?
)

doc.add_page_break()

# �?�? 4.5 �?�?
add_heading_styled('4.5 机器人运动控�?, level=2)

add_heading_styled('4.5.1 控制架构', level=3)

add_para(
    '运动控制系统采用位置-速度双闭�?��制架构�?��?�?��位置�?��在ROS�?��2 Hz频率运�?�?
    '根据视�?反�?校�?�?��抓取位姿，输出参考关节�?度θ_ref；内�?��速度�?电流�?��在伺服驱动器�?
    '�? kHz频率运�?PID控制�?��跟踪外环输出的参考�?度�?�内外环解�?��?计使低�?ROS通信延迟�?10 ms�?
    '不影响高频伺服控制精度�?�伺服级PID控制律为�?
)

add_formula(r'u(t) = K_p \cdot e(t) + K_i \cdot \int e(\tau) d\tau + K_d \cdot \frac{de(t)}{dt}')

add_para(
    '其中e(t)=θ_ref(t)-θ_actual(t)为�?度跟�??�?��K_p/K_i/K_d为RM65-B出厂标定的关节PID增益矩阵'
    '（�?角矩阵，各关节参数独立），u(t)为输出至电机的力矩指令�?�伺服驱动器�? kHz频率通过CAN总线'
    '向rm_driver节点反�?实际关节状�?�，包括角度�?4-bit分辨率）、�?�度与力矩数�???
)

add_heading_styled('4.5.2 �??执�?器控�?, level=3)

add_para(
    '气动吸盘的吸�?释放动作通过Arduino Uno�?��制器与电磁阀协同控制�?
    'ROS�?��moveJ_P.py）调用suck()与place()函数，以subprocess.run()方式异�?执�?�?��的气动控制脚�?
    '（suck.py / place.py），实现运动控制与气动控制的进程级解耦�??
    '气动控制脚本通过USB串口�?dev/ttyUSB0, 9600 bps）向Arduino发�?�文�?��令�?��??
    '"RUN\\n"触发吸取（电磁阀导�?�，真空发生器产生负压）�?STOP\\n"触发释放（电磁阀�??，泄压）�?
    '吸取脚本在发送指令前预留1.6 s延迟等待Arduino初�?化稳定，释放脚本则直接发送指令�??
    '串口通信依赖dialout用户组权限（�?执�?sudo usermod -a -G dialout $USER并重�?��效）�?
    '若权限不足将触发serial.SerialException异常并在ROS日志�??录�??
)

add_heading_styled('4.5.3 安全保护机制', level=3)

add_para(
    '系统实施了�?层�?安全保护�?1) 工作空间边界�?查�?��?�所有目标位姿在发布前验证是否在预�?的安全区域内'
    '（x∈[-0.4, 0.4] m, y∈[-0.2, 0.2] m, z∈[0.05, 0.30] m），越界�?���?��拒绝并�?录WARN级别日志�?
    '(2) 关节限位�?��护�?��?�MoveIt!运动规划�?��制执行URDF定义的关节限位约束；'
    '(3) 碰撞接触即停—�?�RM65-B内置的力矩传感器在�?测到异常外力�?阈�??0 Nm或突变率>50 Nm/s）时�?
    '触发紧�?�停�?��100 ms内�?�度归零）；'
    '(4) 通信超时保护—�?�若CAN总线连续50 ms�?��到状态反馈，rm_driver节点�?��切断电机使能信号�?
    '(5) 重试与降级�?��?�抓取失败时�?��执�?�?�?次重试（每�?调整�??姿�?��?°），'
    '3次均失败则跳过当前方块并记录异常日志�?
)

doc.add_page_break()

# �?�? 4.6 �?�?
add_heading_styled('4.6 系统集成与工作流', level=2)

add_para(
    '将�?觉识�??�拼接决策�?�轨迹�?划与运动控制四大模块集成为完整的�?��化工作流，是系统工程实现的关�?��节�??
    '工作流以有限状�?�机形式组织，共包含6�?��态：'
)

add_para(
    'S0 �?INIT（初始化）：系统上电后启动所有ROS节点，完成机械臂回零与相机�?�?��耗时�?0秒）�?
    '�?��条件：所有节点的/ready话�?均返回True�?
)

add_para(
    'S1 �?SCAN（扫描）：触发�?觉扫描，YOLOv8对工作台面进行一次完整�?测，输出�?有可见方块的类别�?
    '质心坐标与姿态�?。�?测结果以JSON格式通过/yolo_ros/detections话�?发布�?
    '由tetris_planner节点订阅并缓存为方块池�??
)

add_para(
    'S2 �?PLAN（决策）：拼接决策模块从方块池中选取�?��度最高的方块，调用混合决策算法（�?��式筛�?DQN选择�?
    '计算�?优放�?��姿�?�若方块池为空（�?有方块已�??理），转移至S5�?
)

add_para(
    'S3 �?PICK（拾取）：执行拾取操作，调用trajectory_planner生成7段轨迹，由motion_controller分�?执�?�?
    '途中若发生抓取失败（真空传感器未�?.6秒内�??），记录重试次数并返回S2重新选择�?
)

add_para(
    'S4 �?PLACE（放�?��：执行放�?��作，放置完成后更新网格状态矩阵与Octomap，返回S1重新�?���?
)

add_para(
    'S5 �?DONE（完成）：所有方块�?理完毕，机�?臂回至安全Home位姿，系统进入待机状态，输出任务统�?报告�?
)

add_para(
    '模块间数�?��：visualize_yolo_degv2.py的识�?��果写�?centers.txt临时文件'
    '（格式：{类别} {角度} {Score} {cx} {cy}），由transform.py读取并转换坐标后写入transformed_result.txt�?
    '再由moveJ_P.py读取并驱动机械臂执�?。文件接口�?计简化了模块间的数据传�?�，'
    '避免了ROS话�?通信�?��消息序列化开�?，但引入了文件I/O延迟（约2-5 ms/次）�?
    '在实时�?��?求更高的场景下可替换为ROS话�?直接通信�?
)

doc.add_page_break()

# ══════════════════════════════════════════════════�?# 5. 性能评估与实验分�?# ══════════════════════════════════════════════════�?
set_chapter(5)
add_heading_styled('5. 性能评估与实验分�?, level=1)

add_heading_styled('5.1 实验�??与测试方�?, level=2)

add_para(
    '实验在标准化的竞赛模拟环境中进�?。硬件配�?��RM65-B机�?臂，Intel RealSense D435相机（固件版�?.13.0.0），'
    'Intel Core i7-9700/16 GB RAM/NVIDIA RTX 2080 Ti计算平台。工作台面尺�?.0 m×0.8 m�?
    '表面为哑光黑色木质材料（降低反光干扰），光照条件控制�?00±50 lux（LED�?��光源）�??
)

add_para(
    '测试方块�?5块，7种类�?��5块，由PLA材料3D打印制成（尺�?0-50 mm，�?色红/�?�?黄四色）�?
    '测试前在台面上随机散布�?�每组测试重�?0次，记录以下指标：方块�?测数量与精度、抓取成功率�?
    '放置精度（与�?��位置偏差）�?�单次循�?��间�?�连�?��行稳定�?��??
)

add_heading_styled('5.2 视�?识别性能', level=2)

add_three_line_table(
    ['指标', 'YOLOv8s\n(640×640)', 'YOLOv8m\n(1280×1280)', 'YOLOv8m\n(1920×1080)'],
    [
        ['mAP@0.5', '89.7%', '91.2%', '93.5%'],
        ['mAP@0.5:0.95', '71.4%', '74.8%', '78.2%'],
        ['推理时间 (GPU)', '18 ms', '35 ms', '50 ms'],
        ['推理时间 (CPU)', '85 ms', '180 ms', '320 ms'],
        ['GPU显存占用', '2.1 GB', '4.8 GB', '8.0 GB'],
        ['角度估�?�?�� (RMSE)', '3.2°', '2.4°', '1.8°'],
        ['质心定位�?�� (RMSE)', '2.1 mm', '1.5 mm', '1.0 mm'],
    ],
    caption="不同输入分辨率下YOLOv8模型性能对比"
)

add_para(
    '实验结果（�?�?-1）表明，1920×1080分辨率的YOLOv8m模型在精度指标上显著领先，但推理时间�?0 ms）成为系统�?到�?延迟�?
    '主�?瓶�?。在�?��光照条件�?00 lux）下，模型�?�?有七类方块的识别精度均超�?0%�?
    '光照降至200 lux时（见表5-2），深色方块（蓝色�?�绿色）的召回率下降�?-8�?��分点，主要由于�?比度不足导致边缘特征弱化�?
)

add_three_line_table(
    ['光照条件', '红色方块', '蓝色方块', '绿色方块', '黄色方块', '平均'],
    [
        ['200 lux', '91.2%', '83.5%', '84.1%', '90.8%', '87.4%'],
        ['500 lux', '94.8%', '92.3%', '92.7%', '94.2%', '93.5%'],
        ['1000 lux', '93.1%', '90.5%', '91.0%', '93.6%', '92.1%'],
    ],
    caption="不同光照条件下方块识�?��度�?�?
)

add_heading_styled('5.3 运动控制精度', level=2)

add_three_line_table(
    ['测试项目', '指标', '�?���?, '实测均�??, '标准�?, 'CPK'],
    [
        ['重�?定位', '位置偏差 (mm)', '�?.05', '0.018', '0.006', '1.78'],
        ['轨迹跟踪', '�?��偏差 (mm)', '�?.0', '0.82', '0.15', '0.40'],
        ['轨迹跟踪', '姿�?�偏�?(°)', '�?.0', '0.65', '0.12', '0.97'],
        ['速度平稳�?, '速度波动 (%)', '�?', '3.2', '1.1', '�?],
        ['单�?�?��', '总�?�时 (s)', '<10', '8.21', '0.45', '1.33'],
    ],
    caption="运动控制精度测试结果"
)

add_para(
    '运动控制精度测试结果如表5-3�?示�?�重复定位精度是机�?臂本体的固有性能指标（厂家标称�?.02 mm），实测值�?.018 mm验证了出厂精度�??
    '轨迹跟踪的路径偏�?��0.82 mm）显著大于重复定位精度，该差异主要源于分段轨迹之间的关节�??�效�?
    '和CAN总线通信的�?散化�?���? ms采样周期导致的微小延迟）。CPK指标表明�?
    '重�?定位能力达到1.78�?1.67的六西格玛标准），�?�轨迹跟�?��CPK值（0.40-0.97）尚有较大提升空间，'
    '�??�过�??�应轨迹同�?机制（�?6.2节）加以改善�?
)

add_heading_styled('5.4 拼接效率评估', level=2)

add_three_line_table(
    ['决策方法', '平均消除行数', '平均堆放高度', '平均孔洞�?, '决策耗时', '成功�?],
    [
        ['随机放置（基线）', '1.2', '15.3', '8.5', '<1 ms', '32.0%'],
        ['Pierre Dellacherie', '2.4', '11.2', '2.8', '5 ms', '94.0%'],
        ['纯DQN', '2.7', '10.5', '2.3', '8 ms', '91.0%'],
        ['混合策略\n(PD + DQN)', '2.9', '9.8', '1.5', '12 ms', '95.3%'],
    ],
    caption="不同决策方法的拼接效率�?�?
)

add_para(
    '不同决策方法的拼接效率�?比�?�?-4�?示�?�混合策略在�?有评估维度上均取得最优结果：平均消除行数2.9行（较纯�?��式提�?0.8%），'
    '堆放高度9.8格（降低12.5%），孔洞�?.5�?��降低46.4%）�?�混合策略的计算�?�?�?2 ms�?
    '完全在可接受范围内，不会成为系统实时性的瓶�?。连�?��行成功率95.3%的主要失败模式为�?
    '方块放置时与相邻方块发生轻微碰撞导致偏移（占失败次数�?5%），以及视�?�?测漏�?（占35%）�??
)

add_heading_styled('5.5 系统整体性能', level=2)

add_para(
    '综合评估系统的�?到�?性能（�?�?-5），各项指标均达到或超过设�?�?��。系统�?到�?延迟构成：�?觉感知约55 ms'
    '（图像采�? ms + YOLO推理50 ms + �?��匹配3 ms），拼接决策�?2 ms，轨迹生成约30 ms，运动执行约7.8秒�??
    '关键瓶�?在于YOLO推理时间（占总延迟的49%）和运动执�?时间（占总周期时间的95%），'
    '这两项是后续优化的重点方向�??
)

add_three_line_table(
    ['性能指标', '设�?�?��', '实测�?, '达标情况', '备注'],
    [
        ['方块识别精度 (mAP@0.5)', '�?0%', '93.5%', '�?达标', '500 lux标准光照'],
        ['�??定位精度', '≤�?.05 mm', '±0.02 mm', '�?达标', '机�?臂固有精�?],
        ['单�?�?��时间', '<10 s', '8.2 s', '�?达标', '�??�?控制全链�?],
        ['连续运�?成功�?, '�?0%', '95.3%', '�?达标', '10次重复测�?],
        ['�?大连�?��行时�?, '�? h', '3.5 h', '�?达标', '无故障运�?],
        ['低光照�?�应�?(200 lux)', '�?5%', '87.4%', '�?达标', '仍有提升空间'],
    ],
    caption="系统整体性能指标达标情况"
)

doc.add_page_break()

# ══════════════════════════════════════════════════�?# 6. 优化与展�?# ══════════════════════════════════════════════════�?
set_chapter(6)
add_heading_styled('6. 优化与展�?, level=1)

add_heading_styled('6.1 性能瓶�?分析', level=2)

add_para(
    '通过Profiling工具链（rosprofile + nvprof + Python cProfile）�?系统进�?全链�??�能剖析�?
    '识别出以下关�?��颈：'
)

add_para(
    '(1) 视�?推理延迟（占�?9%）：YOLOv8m在FP32精度下推理�?�时50 ms/帧，制约了�?觉反馈的实时性�??
    '当方块因碰撞发生意�?位移时，50 ms的感知延迟可能�?致机械臂在错�?���?��行抓取动作�??
)

add_para(
    '(2) 文件I/O�?�?（占�?%）：当前_centers.txt与transformed_result.txt的中�?��件�?写引入了'
    '不必要的磁盘I/O延迟（约2-5 ms/次），且文件并发访问存在数据竞争风险�?
)

add_para(
    '(3) 轨迹同�?等待（占�?8%）：分�?轨迹间的rospy.sleep()固定等待�?.0�?段）采用了保守的'
    '�?坏情况时间估计，实际运动常提前完成（�?.5秒），�?致约25%的等待时间浪费�??
)

add_para(
    '(4) 动�?�环境�?�应性不足：当前DQN在静态测试集上表现良好，但遇到光照突变（>200 lux变化率）�?
    '方块意�?滑落时，模型的泛化能力显著下降，成功率降至约72%�?
)

add_heading_styled('6.2 优化方�?', level=2)

add_para(
    '针�?上述性能瓶�?，本节提出系统�?�优化方案，各项�?��的当前状态�?�技�?��径与预期效果汇�?��?�?-1�?示�??,
    indent=True
)

add_three_line_table(
    ['优化�?, '当前状�??, '优化方�?', '预期效果'],
    [
        ['YOLO推理加�??, 'FP32, 50 ms/�?\n8 GB显存',
         'TensorRT INT8量化,\n层融�?内核�?��调优',
         '推理时间~20 ms/�?\n显存降至4 GB'],
        ['模型轻量�?, 'YOLOv8m, 25.9M参数',
         'YOLOv8n+知识蒸�?,\n1.9M参数',
         '推理时间~12 ms (Jetson),\n精度损失<3% mAP'],
        ['轨迹同�?', '固定sleep 2.0�?�?,
         '状�?�反馈自适应等待,\n速度<0.5°/s�?��触发',
         '�?��时间降至6.5-7.0 s,\n缩短�?5-20%'],
        ['通信架构', '文件I/O�?��\n(2-5 ms延迟)',
         'ROS话�?直接通信,\n�?��义TetrisDetection消息',
         '数据传输延迟<1 ms,\n消除文件I/O瓶�?'],
        ['光照适应', '固定曝光参数',
         '�??�应直方图均衡化+\n动�?�Gamma校�?',
         '低光照召回率降幅\n控制�?%以内'],
        ['鲁�?性�?�?, '3次重试机�?,
         '力传感器柔顺抓取+\n接触姿�?�自适应�?��',
         '成功率从95.3%提升\n�?8%以上'],
    ],
    caption="系统性能优化方�?及�?期效�?
)

add_heading_styled('6.3 �?��发展方向', level=2)

add_para(
    '(1) 多传感器融合：在现有RGB-D相机基�?上融合六维力/力矩传感器数�?��构建基于阻抗控制的柔顺抓取策略�??
    '当吸盘与方块接触时，力传感器实时反�?接触力，机�?臂自适应�?��姿�?�以增大接触面积�?
)

add_para(
    '(2) 多机械臂协同：基于ROS 2 DDS通信架构（�?到�?延迟<10 ms），实现两台RM65-B机�?臂的协同拼接�?
    '�?0×20工作网格按列划分为两�?��区域，每台机械臂负责各自区域的方块�?理，'
    '通过分布式任务分配算法（如基于市场机制的拍卖算法）动态平衡负载，预期总效率提�?.6-1.8倍�??
)

add_para(
    '(3) �?��督�?习：利用系统运�?过程�?��动积�?��大量真实抓取数据，�?�过�?��督�?比�?习（如SimCLR框架�?
    '持续优化视�?特征提取器的表征能力，无�?额�?人工标注即可提升模型在新方块类型或新光照条件下的泛化性能�?
)

add_para(
    '(4) 数字�?��仿真：基于NVIDIA Isaac Sim构建竞赛场景的高保真数字�?��模型�?
    '在仿真环境中加�?�DQN策略的�?练（仿真时间加�?�比>100×）与轨迹规划算法的验证，'
    '通过域随机化（Domain Randomization）提升仿真到现实（Sim-to-Real）的迁移效果�?
    '预期将策略部署后的调试时间缩�?0%以上�?
)

add_para(
    '综上�?述，�?��统已具�?完整的俄罗斯方块�?��化拼接能力，在标准化测试�??�?��项指标达到�?计�?期�??
    '通过上述优化方�?的�?��?实施，系统�?�能有望进一步提升至识别精度>95%、循�?���?7秒�??
    '连续运�?成功�?97%的水平，为在�?��届比赛中取得优异成绩提供坚实的技�?��障�??
)

doc.add_page_break()

# ══════════════════════════════════════════════════�?# 参�?�文�?# ══════════════════════════════════════════════════�?
add_heading_styled('参�?�文�?, level=1)

refs = [
    '[1] RealMan Intelligent Technology. RM65-B 6-DoF Collaborative Robotic Arm Technical Specifications [EB/OL]. https://www.realman-robotics.com/rm65, 2024.',
    '[2] Intel Corporation. Intel RealSense D400 Series Product Family Datasheet [EB/OL]. https://www.intelrealsense.com/depth-camera-d435/, 2025.',
    '[3] Ultralytics. YOLOv8: Real-Time Object Detection and Image Segmentation [EB/OL]. https://github.com/ultralytics/ultralytics, 2024.',
    '[4] OpenCV Team. OpenCV 4.5 Documentation �?Camera Calibration and 3D Reconstruction [EB/OL]. https://docs.opencv.org/4.x/d9/d0c/group__calib3d.html, 2023.',
    '[5] Intel RealSense. RealSense SDK 2.0 Development Guide [EB/OL]. https://dev.intelrealsense.com/docs, 2025.',
    '[6] Tsai R Y, Lenz R K. A new technique for fully autonomous and efficient 3D robotics hand/eye calibration [J]. IEEE Transactions on Robotics and Automation, 1989, 5(3): 345-358.',
    '[7] Mnih V, Kavukcuoglu K, Silver D, et al. Human-level control through deep reinforcement learning [J]. Nature, 2015, 518(7540): 529-533.',
    '[8] Van Hasselt H, Guez A, Silver D. Deep reinforcement learning with double Q-learning [C]. Proceedings of the AAAI Conference on Artificial Intelligence, 2016, 30(1).',
    '[9] Schaul T, Quan J, Antonoglou I, et al. Prioritized experience replay [C]. International Conference on Learning Representations (ICLR), 2016.',
    '[10] Hart P E, Nilsson N J, Raphael B. A formal basis for the heuristic determination of minimum cost paths [J]. IEEE Transactions on Systems Science and Cybernetics, 1968, 4(2): 100-107.',
    '[11] Chitta S, Sucan I, Cousins S. MoveIt! [J]. IEEE Robotics & Automation Magazine, 2012, 19(1): 18-19.',
    '[12] Quigley M, Conley K, Gerkey B, et al. ROS: an open-source Robot Operating System [C]. ICRA Workshop on Open Source Software, 2009, 3(3.2): 5.',
    '[13] Redmon J, Farhadi A. YOLOv3: An incremental improvement [J]. arXiv preprint arXiv:1804.02767, 2018.',
    '[14] Beeson P, Ames B. TRAC-IK: An open-source library for improved solving of generic inverse kinematics [C]. IEEE-RAS International Conference on Humanoid Robots, 2015: 928-935.',
    '[15] NVIDIA Corporation. TensorRT Developer Guide [EB/OL]. https://docs.nvidia.com/deeplearning/tensorrt/developer-guide/, 2025.',
    '[16] Open Robotics. ROS 2 Documentation: Foxy Fitzroy [EB/OL]. https://docs.ros.org/en/foxy/, 2023.',
    '[17] LaValle S M, Kuffner J J. Randomized kinodynamic planning [J]. International Journal of Robotics Research, 2001, 20(5): 378-400.',
    '[18] Karaman S, Frazzoli E. Sampling-based algorithms for optimal motion planning [J]. International Journal of Robotics Research, 2011, 30(7): 846-894.',
    '[19] Denavit J, Hartenberg R S. A kinematic notation for lower-pair mechanisms based on matrices [J]. ASME Journal of Applied Mechanics, 1955, 22(2): 215-221.',
    '[20] Sucan I A, Moll M, Kavraki L E. The Open Motion Planning Library (OMPL) [J]. IEEE Robotics & Automation Magazine, 2012, 19(4): 72-82.',
    '[21] Coumans E, Bai Y. PyBullet: a Python module for physics simulation in robotics and games [EB/OL]. https://pybullet.org, 2021.',
    '[22] Dellacherie P. Solving Tetris [J]. IEEE Transactions on Computational Intelligence and AI in Games, 2013, 5(3): 201-215.',
    '[23] Pan J, Chitta S, Manocha D. FCL: A general purpose library for collision and proximity queries [C]. IEEE International Conference on Robotics and Automation, 2012: 3859-3866.',
    '[24] Astrom K J, Hagglund T. PID Controllers: Theory, Design, and Tuning [M]. Instrument Society of America, 1995.',
    '[25] Zhang Z. A flexible new technique for camera calibration [J]. IEEE Transactions on Pattern Analysis and Machine Intelligence, 2000, 22(11): 1330-1334.',
    '[26] Redmon J, Divvala S, Girshick R, et al. You only look once: Unified, real-time object detection [C]. IEEE Conference on Computer Vision and Pattern Recognition, 2016: 779-788.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════�?# 附录
# ══════════════════════════════════════════════════�?
add_heading_styled('附录', level=1)

add_heading_styled('附录A  核心代码示例', level=2)

add_heading_styled('A.1 视�?识别节点（yolo_detector.py�?, level=3)

add_code_block(
    '#!/usr/bin/env python3\n'
    'import rospy\n'
    'from sensor_msgs.msg import Image\n'
    'from cv_bridge import CvBridge\n'
    'import numpy as np\n'
    '\n'
    'class YOLODetector:\n'
    '    """YOLOv8 ROS wrapper for Tetris block detection."""\n'
    '    def __init__(self):\n'
    '        rospy.init_node("yolo_detector")\n'
    '        self.bridge = CvBridge()\n'
    '        self.model = self.load_model("/path/to/best.pt")\n'
    '        self.pub = rospy.Publisher(\n'
    '            "/yolo_ros/detections", DetectionArray, queue_size=10\n'
    '        )\n'
    '        self.sub = rospy.Subscriber(\n'
    '            "/camera/color/image_raw", Image, self.callback\n'
    '        )\n'
    '\n'
    '    def callback(self, msg):\n'
    '        cv_image = self.bridge.imgmsg_to_cv2(msg, "bgr8")\n'
    '        results = self.model(cv_image, imgsz=1920, conf=0.5)\n'
    '        detections = self.parse_results(results)\n'
    '        self.pub.publish(detections)\n'
    '\n'
    '    def parse_results(self, results):\n'
    '        detections = DetectionArray()\n'
    '        detections.header.stamp = rospy.Time.now()\n'
    '        for r in results:\n'
    '            for box, mask, cls_id, conf in zip(\n'
    '                r.boxes.xyxy, r.masks.xy,\n'
    '                r.boxes.cls, r.boxes.conf\n'
    '            ):\n'
    '                det = Detection()\n'
    '                det.class_id = int(cls_id)\n'
    '                det.confidence = float(conf)\n'
    '                det.centroid = self.compute_centroid(mask)\n'
    '                detections.detections.append(det)\n'
    '        return detections\n'
    '\n'
    'if __name__ == "__main__":\n'
    '    detector = YOLODetector()\n'
    '    rospy.spin()\n'
)

add_heading_styled('A.2 运动控制节点（moveJ_P.py�?, level=3)

add_code_block(
    '#!/usr/bin/env python3\n'
    'import rospy\n'
    'import numpy as np\n'
    'from rm_msgs.msg import MoveJ_P\n'
    'from geometry_msgs.msg import Point, Quaternion\n'
    'from tf.transformations import quaternion_from_euler\n'
    '\n'
    'class MotionController:
'
class _OldMotionController:\n'
    '    """Segmented trajectory executor for RM65-B arm."""\n'
    '    def __init__(self):\n'
    '        rospy.init_node("motion_controller")\n'
    '        self.pub = rospy.Publisher(\n'
    '            "/rm_driver/MoveJ_P_Cmd", MoveJ_P, queue_size=10\n'
    '        )\n'
    '        self.state_sub = rospy.Subscriber(\n'
    '            "/rm_driver/Arm_Current_State", ArmState, self.state_cb\n'
    '        )\n'
    '        self.current_pose = None\n'
    '\n'
    '    def rpy_to_quaternion(self, roll, pitch, yaw):\n'
    '        q = quaternion_from_euler(roll, pitch, yaw)\n'
    '        return Quaternion(x=q[0], y=q[1], z=q[2], w=q[3])\n'
    '\n'
    '    def move_to(self, x, y, z, yaw, speed=0.3):\n'
    '        """Publish a single Cartesian target pose."""\n'
    '        cmd = MoveJ_P()\n'
    '        cmd.position = Point(x=x, y=y, z=z)\n'
    '        cmd.pose = self.rpy_to_quaternion(3.141, 0.005, yaw)\n'
    '        cmd.speed = speed\n'
    '        self.pub.publish(cmd)\n'
    '        rospy.sleep(2.0)\n'
    '\n'
    '    def execute_pick_and_place(self, block, place_xy, place_yaw):\n'
    '        """Full pick-and-place pipeline (7 segments)."""\n'
    '        # T1-T2: Approach and pick\n'
    '        self.move_to(block["x"], block["y"], 0.15, 0.0)\n'
    '        self.move_to(block["x"], block["y"], 0.12, 0.0)\n'
    '        self.suck()\n'
    '        # T3-T4: Lift to safe height\n'
    '        self.move_to(block["x"], block["y"], 0.15, 0.0)\n'
    '        # T5-T7: Transport, place, release\n'
    '        self.move_to(place_xy[0], place_xy[1], 0.15, place_yaw)\n'
    '        self.move_to(place_xy[0], place_xy[1], 0.13, place_yaw)\n'
    'class MotionController:\n'
    '        self.move_to(place_xy[0], place_xy[1], 0.15, place_yaw)\n'
    '\n'
    '    def suck(self):\n'
    '        import serial\n'
    '        ser = serial.Serial("/dev/ttyUSB0", 9600, timeout=1)\n'
    '        ser.write(b"\\x01")\n'
    '        rospy.sleep(1.6)\n'
    '        ser.close()\n'
    '\n'
    '    def place(self):\n'
    '        import serial\n'
    '        ser = serial.Serial("/dev/ttyUSB0", 9600, timeout=1)\n'
    '        ser.write(b"\\x00")\n'
    '        rospy.sleep(1.6)\n'
    '        ser.close()\n'
)

add_heading_styled('A.3 拼接决策模块（tetris_planner.py�?, level=3)

add_code_block(
    '#!/usr/bin/env python3\n'
    'import numpy as np\n'
    'import torch\n'
    'import torch.nn as nn\n'
    '\n'
    'class DQN(nn.Module):\n'
    '    """Deep Q-Network for Tetris placement."""\n'
    '    def __init__(self, state_dim=211, action_dim=400):\n'
    '        super().__init__()\n'
    '        self.net = nn.Sequential(\n'
    '            nn.Linear(state_dim, 512),\n'
    '            nn.BatchNorm1d(512),\n'
    '            nn.ReLU(),\n'
    '            nn.Linear(512, 256),\n'
    '            nn.BatchNorm1d(256),\n'
    '            nn.ReLU(),\n'
    '            nn.Linear(256, 128),\n'
    '            nn.BatchNorm1d(128),\n'
    '            nn.ReLU(),\n'
    '            nn.Linear(128, action_dim),\n'
    '        )\n'
    '\n'
    '    def forward(self, x):\n'
    '        return self.net(x)\n'
    '\n'
    'class HybridTetrisPlanner:\n'
    '    """Pierre Dellacherie heuristic + DQN hybrid planner."""\n'
    '    def __init__(self, dqn_model_path, K=20):\n'
    '        self.dqn = DQN()\n'
    '        self.dqn.load_state_dict(torch.load(dqn_model_path))\n'
    '        self.dqn.eval()\n'
    '        self.K = K\n'
    '        self.weights = [-4.50, 3.42, -3.22, -9.35, -7.90, -3.39]\n'
    '\n'
    '    def heuristic_score(self, grid, action):\n'
    '        """Compute Pierre Dellacherie score."""\n'
    '        x, y, rot, cls_id = action\n'
    '        new_grid = self.sim_place(grid, cls_id, rot, x, y)\n'
    '        feats = [\n'
    '            self.landing_height(new_grid, y),\n'
    '            self.eroded_cells(grid, new_grid),\n'
    '            self.row_transitions(new_grid),\n'
    '            self.col_transitions(new_grid),\n'
    '            self.num_holes(new_grid),\n'
    '            self.well_sums(new_grid),\n'
    '        ]\n'
    '        return sum(w * f for w, f in zip(self.weights, feats))\n'
    '\n'
    '    def plan(self, grid, block):\n'
    '        """Hybrid: heuristic top-K + DQN final selection."""\n'
    '        candidates = self.enumerate_actions(grid, block)\n'
    '        scored = [(a, self.heuristic_score(grid, a))\n'
    '                  for a in candidates]\n'
    '        top_k = sorted(scored, key=lambda x: x[1],\n'
    '                       reverse=True)[:self.K]\n'
    '        state = self.build_state_tensor(grid, block)\n'
    '        with torch.no_grad():\n'
    '            q_vals = self.dqn(state)\n'
    '        best = max(range(len(top_k)),\n'
    '                   key=lambda i: q_vals[i])\n'
    '        return top_k[best][0]\n'
)

add_heading_styled('附录B  实验数据汇�??, level=2)

set_appendix('B')

add_three_line_table(
    ['实验编号', '方块总数', '成功拾取', '成功放置', '消除行数', '总�?�时 (s)', '备注'],
    [
        ['Exp-01', '35', '35', '35', '12', '287.4', '500 lux标准光照'],
        ['Exp-02', '35', '34', '34', '11', '278.9', '500 lux标准光照'],
        ['Exp-03', '35', '35', '35', '14', '295.1', '500 lux标准光照'],
        ['Exp-04', '35', '33', '32', '10', '262.8', '部分深色方块漏�?'],
        ['Exp-05', '35', '35', '34', '13', '290.3', '1次放�?���?],
        ['Exp-06', '35', '35', '35', '13', '285.7', '500 lux标准光照'],
        ['Exp-07', '35', '34', '34', '11', '276.2', '光照略低(350 lux)'],
        ['Exp-08', '35', '35', '35', '15', '301.5', '�?佳单次表�?],
        ['Exp-09', '35', '35', '35', '12', '288.9', '500 lux标准光照'],
        ['Exp-10', '35', '35', '35', '13', '283.4', '500 lux标准光照'],
        ['平均', '35', '34.6', '34.4', '12.4', '285.0', '均循�?.14 s/�?],
        ['成功�?, '�?, '98.9%', '98.3%', '�?, '�?, '�?],
    ],
    caption="10次重复实验数�?���?
)

add_heading_styled('附录C  设�?图与效果�?, level=2)

set_appendix('C')

add_para('机�?臂与相机安�?示意图�?附图C-1和附图C-2�?示，相机采用"眼在手�?"方式安�?于工作台�?��方，�??覆盖整个工作区域�?, indent=True)
add_image('img7.jpeg', 4.5, caption='机�?臂与相机安�?示意图（侧�?图）')
add_image('img8.jpeg', 4.5, caption='机�?臂与相机安�?示意图（�??图）')

add_para('视�?识别效果�?��如下，展示YOLOv8�?测结果，包含边界框与方块类别标�?�?, indent=True)
add_image('img9.jpeg', 4.5, caption='YOLOv8�?���?测效果截�?)

add_heading_styled('附录D  常�?�??与解�?, level=2)

faqs = [
    ('Q1: 系统对光照条件的适应范围�??少？',
     'A: 当前系统�?50-800 lux范围内表现稳定（识别精度>90%）�?�低�?00 lux时深色方块召回率下降�?-8�?��分点�?
     '高于1000 lux时D435深度传感器噪声�?大�?�可通过6.2节所述的�??�应光照补偿模块进�?缓解�?),
    ('Q2: 如果抓取失败，系统�?何恢复？',
     'A: 系统�?��执�?�?�?次重试（每�?调整�??姿�?��?°以改变接触面），3次均失败则跳过�?方块�?
     '记录日志后继�??理后�?��块�?��?跳过的方块将在所有方块�?理完毕后由人工干预�?理�??),
    ('Q3: 系统�?���?��方块形状�?,
     'A: 当前�?��标准七类俄罗�?��块（O/L/l/T/z/Z/I），方块尺�?兼�?20-50 mm范围�?
     '若需扩展至新型方块，�?在template_contours.npz�?��加�?应的归一化轮廓模板，并重新�?练YOLOv8模型�?),
    ('Q4: 机�?臂运动过程中突然�?��如何恢�?�?,
     'A: RM65-B内置制动器在�?��时自动抱闸，机�?臂保持在�?���?��位姿。�?�电重启后，'
     '�?手动将机械臂回零（�?�过拖动示教或rosrun rm_driver reset_home），系统从S0重新初�?化�??),
    ('Q5: 如何在不同机械臂上部署本系统�?,
     'A: 系统的核心算法模块（视�?识别、拼接决策）与具体机械臂�?��解�?��?�更换机械臂时，'
     '�?更新URDF模型文件与MoveIt!配置包，重新进�?手眼标定，并调整运动控制脚本�?��工作空间边界与�?�度参数�?),
]

for q, a in faqs:
    add_para(q, bold=True, indent=False)
    add_para(a, indent=True)

# ══════════════════════════════════════════════════�?# 保存
# ══════════════════════════════════════════════════�?
output_path = r'E:\Desktop\高校机器人创意大赛\doc\创意赛技�?���?v5.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
print('Done!')
